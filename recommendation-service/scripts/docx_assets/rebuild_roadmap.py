# -*- coding: utf-8 -*-
r"""Viết lại toàn bộ AI_Recommendation_Roadmap.docx cho mạch lạc, gọn hơn.

Tài liệu gốc được viết dần theo 4 phase (Phần I-IV), mỗi phần lặp lại gần như nguyên vẹn
phần giải thích kiến trúc/nguyên tắc thiết kế/thuật toán của phần trước, và 3 lần phát hiện
cùng 1 lớp bug (thiếu lọc active/quantity>0) được kể lại thành 3 đoạn tách rời. Script này
dựng lại cấu trúc mới:

  PHẦN I   — Tổng quan & kiến trúc (nguyên tắc thiết kế viết 1 lần duy nhất)
  PHẦN II  — Các thuật toán/kỹ thuật (giải thích 1 lần, có mô tả ngắn cho cả các thuật toán
             trước đây chỉ nhắc tên: Apriori, DPP, Thompson Sampling)
  PHẦN III — Pipeline hiện tại (bức tranh cuối cùng sau khi gộp cả 4 phase, không lặp lại
             theo từng phase nữa) + vòng lặp đo lường continual learning
  PHẦN IV  — Lịch sử triển khai theo phase (rút gọn, tham chiếu ngược lại Phần II thay vì
             giải thích lại thuật toán)
  PHẦN V   — Sự cố đã phát hiện & bài học (gộp 2 lớp bug lặp lại nhiều lần thành 1 mục mỗi
             loại thay vì kể 3 lần rải rác, bổ sung phát hiện mới nhất từ code review)
  PHẦN VI  — Kiểm chứng & vận hành hiện tại

Toàn bộ số liệu/pseudocode đã đối chiếu lại với code thật trong repo tại thời điểm viết lại
(xem PHẦN V mục cuối để biết những gì vừa được sửa). 6 sơ đồ matplotlib gốc được tái sử dụng
nguyên vẹn (trích xuất từ file docx cũ), chỉ đánh số lại cho đúng thứ tự xuất hiện thật.

Chạy: cd recommendation-service && venv\Scripts\python scripts/docx_assets/rebuild_roadmap.py
(Đóng file Word trước khi chạy — script ghi đè trực tiếp AI_Recommendation_Roadmap.docx.)
"""
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor

ROOT = Path(__file__).resolve().parents[2]
DOCX = ROOT / "AI_Recommendation_Roadmap.docx"
ASSETS = Path(__file__).resolve().parent

# ---------------------------------------------------------------- docx helpers

def title(d, t):
    d.add_paragraph(t, style="Title")


def h1(d, t):
    d.add_paragraph(t, style="Heading 1")


def h2(d, t):
    d.add_paragraph(t, style="Heading 2")


def h3(d, t):
    d.add_paragraph(t, style="Heading 3")


def p(d, t):
    d.add_paragraph(t, style="Normal")


def bullet(d, t):
    d.add_paragraph(t, style="List Bullet")


def num(d, t):
    d.add_paragraph(t, style="List Number")


def picture(d, img: Path, caption: str):
    d.add_picture(str(img), width=Inches(6.3))
    d.add_paragraph(caption, style="Caption")


def code_box(d, text: str):
    """1x1 table nền xanh nhạt viền mảnh, font Consolas — đúng style code-box đã dùng
    trong tài liệu gốc (nền F3F6F4, viền C7D3CB, chữ 2A3831 9pt)."""
    table = d.add_table(rows=1, cols=1)
    table.style = "Normal Table"
    cell = table.rows[0].cells[0]

    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F3F6F4")
    tcPr.append(shd)
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "C7D3CB")
        borders.append(el)
    tcPr.append(borders)

    para = cell.paragraphs[0]
    lines = text.strip("\n").split("\n")
    for i, line in enumerate(lines):
        run = para.add_run(line if i == 0 else "")
        if i > 0:
            run.add_break()
            run.text = line
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x2A, 0x38, 0x31)
    d.add_paragraph()  # khoảng trống nhỏ sau code box


def data_table(d, headers, rows):
    table = d.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, htext in enumerate(headers):
        table.rows[0].cells[i].text = htext
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    d.add_paragraph()


# ---------------------------------------------------------------- nội dung

def build(d):
    IMG_ARCH = ASSETS / "img_phase1_architecture.png"
    IMG_PIPE1 = ASSETS / "img_phase1_pipeline.png"
    IMG_ALS = ASSETS / "img_phase2_als.png"
    IMG_PIPE2 = ASSETS / "img_phase2_pipeline.png"
    IMG_LINUCB = ASSETS / "phase3_linucb.png"
    IMG_BANDIT_PIPE = ASSETS / "phase3_pipeline.png"

    title(d, "Hệ thống gợi ý sản phẩm tự học (AI Recommendation System)")
    p(d, "Tài liệu kỹ thuật cho hạng mục ưu tiên cao nhất của dự án webbachhoa: một hệ "
         "thống gợi ý sản phẩm cá nhân hoá, tự học liên tục (continual learning), thay thế "
         "hoàn toàn tính năng \"sản phẩm tương tự\" cũ đã gãy. Cả 4 phase trong roadmap ban "
         "đầu (Phase 0-3) đã triển khai xong và được kiểm chứng qua HTTP/UI thật; tài liệu "
         "này mô tả kiến trúc, thuật toán, và trạng thái vận hành hiện tại — không còn là "
         "một bản kế hoạch chưa làm.")

    # ============================================================ PHẦN I
    h1(d, "PHẦN I — TỔNG QUAN & KIẾN TRÚC")

    h2(d, "1. Bối cảnh & mục tiêu")
    p(d, "Tính năng gợi ý cũ (\"Sản phẩm tương tự\" trên trang chi tiết sản phẩm) đã chết "
         "hoàn toàn — 2 nỗ lực trước đó đều gọi vào endpoint không còn tồn tại. Thay vì chỉ "
         "vá lại, mục tiêu lần này lớn hơn: một mô hình tự học liên tục, cá nhân hoá theo "
         "hành vi thật (mua, xem, tìm kiếm, mua chung), với yêu cầu bắt buộc là không được "
         "rơi vào \"bong bóng lọc\" — khách mua loại A, B rồi chỉ bị gợi ý mãi A, B gây nhàm "
         "chán.")
    p(d, "Khảo sát ban đầu cho thấy dự án không có hạ tầng thu thập hành vi nào (không "
         "lịch sử xem, không log click/impression) và dữ liệu giao dịch thật cực kỳ thưa "
         "(~4 đơn hàng, ~2 user trong dump seed ban đầu). Vì vậy roadmap phải đi theo mốc dữ "
         "liệu — chạy bằng rule-based/thống kê trước, chỉ bật từng phần ML thật khi đạt "
         "ngưỡng dữ liệu tối thiểu (xem mục 4).")

    h2(d, "2. Kiến trúc hệ thống")
    p(d, "3 tầng: Python microservice độc lập (`recommendation-service/`, FastAPI) tính "
         "toán gợi ý và chỉ ĐỌC MySQL trực tiếp → Java Spring Boot (`server/`) đóng vai trò "
         "cầu nối duy nhất, giữ xác thực JWT tập trung và không bao giờ để lỗi của Python "
         "lộ ra khách hàng → React frontend hiển thị và ghi nhận hành vi.")
    picture(d, IMG_ARCH, "Sơ đồ 1 — Kiến trúc 3 tầng và luồng tương tác giữa các thành phần")

    h2(d, "3. Nguyên tắc thiết kế cốt lõi")
    p(d, "Danh sách dưới đây là các ràng buộc thiết kế xuyên suốt toàn bộ hệ thống, áp dụng "
         "cho mọi phase — được liệt kê một lần duy nhất ở đây, các phần sau chỉ tham chiếu "
         "ngược lại thay vì lặp lại.")
    bullet(d, "Fail-safe nghiêm ngặt: Java gọi Python với timeout ngắn (2-2,5 giây). Bất kỳ "
              "lỗi, timeout, hay kết quả rỗng nào cũng lập tức chuyển sang phương án dự "
              "phòng (rule-based/best-seller chạy hoàn toàn trong Java) — không retry, "
              "không để trang trắng hay lỗi hiển thị cho khách.")
    bullet(d, "Python chỉ đọc, không bao giờ ghi: mọi thao tác ghi dữ liệu nghiệp vụ "
              "(click, conversion, impression...) đều đi qua Java. Đây là bất biến kiến "
              "trúc, được giữ nguyên xuyên suốt cả bandit Phase 3.")
    bullet(d, "\"Im lặng\" khi thiếu dữ liệu: mọi nguồn cần ngưỡng dữ liệu tối thiểu "
              "(co-purchase, collaborative filtering) tự trả rỗng khi chưa đủ, không suy "
              "diễn ẩu, không báo lỗi.")
    bullet(d, "Không suy luận ngược dữ liệu lịch sử: giá, khuyến mãi tại thời điểm giao "
              "dịch phải đọc từ snapshot, không bao giờ đọc lại trạng thái Product hiện "
              "tại (có thể đã đổi/hết hạn).")
    bullet(d, "Tồn kho luôn phải tươi nhất: mọi danh sách gợi ý lọc `active=true AND "
              "quantity>0` ngay từ nguồn ứng viên (Python) LẪN ở bước hydrate cuối cùng "
              "(Java, đọc tồn kho tại đúng thời điểm khách xem trang) — 2 lớp lọc vì "
              "Python chỉ tính lại theo chu kỳ, không cập nhật real-time.")
    bullet(d, "Định danh người dùng chỉ lấy từ JWT: `user_id` gửi sang Python hoặc dùng để "
              "ghi dữ liệu nghiệp vụ (attribution, impression) luôn được Java derive từ "
              "token đăng nhập, không bao giờ tin giá trị client tự gửi lên.")
    bullet(d, "Guest vẫn được theo dõi: sessionId ẩn danh sinh ở frontend, mọi hành vi "
              "trước khi đăng nhập được gộp vào tài khoản thật ngay sau khi login "
              "(session-merge).")
    bullet(d, "Không thêm hạ tầng nặng ở quy mô hiện tại: không message broker/Airflow — "
              "polling bảng theo chu kỳ vài phút là đủ; mô hình học lại theo lịch "
              "(nightly cron), không phải theo từng sự kiện (trừ bandit ở Phase 3, xem "
              "Phần II mục 9).")

    h2(d, "4. Dữ liệu & tín hiệu sử dụng")
    p(d, "Toàn bộ tín hiệu bên dưới đọc trực tiếp từ dữ liệu sẵn có của hệ thống (đơn hàng, "
         "sản phẩm) cộng với 4 bảng theo dõi hành vi mới xây ở Phase 0 (`ProductView`, "
         "`SearchLog`, `RecommendationImpression`, `CartEvent`).")
    data_table(
        d,
        ["Nhóm tín hiệu", "Dùng để làm gì"],
        [
            ("Xem sản phẩm (view), tìm kiếm, thêm/bỏ giỏ hàng",
             "Xây hồ sơ sở thích cá nhân hoá trang chủ, tính category-entropy"),
            ("Mua hàng thật (Order + OrderDetail, status PAID/đang giao)",
             "Tín hiệu dương mạnh nhất; input chính cho collaborative filtering"),
            ("Chu kỳ mua lại (khoảng cách trung vị giữa các lần mua lặp)",
             "Đặc thù grocery — hãm tín hiệu \"vừa mua\" (xem Phần II mục 8), giá trị nhất "
             "riêng cho domain bách hóa"),
            ("Khuyến mãi đã snapshot trong đơn (promotionType, freeUnits...)",
             "Đo price sensitivity/promotion responsiveness, dùng được ngay không cần thu "
             "thập thêm"),
            ("Lịch sử hiển thị gợi ý (RecommendationImpression: shown/clicked/converted)",
             "Vòng phản hồi cho continual learning — chống nhàm chán, đo hiệu quả, reward "
             "cho bandit"),
            ("Best-seller & xu hướng 7/30 ngày (Product.sold, OrderDetail GROUP BY)",
             "Lưới an toàn luôn có dữ liệu, kể cả khách vãng lai/sản phẩm mới"),
            ("Mua chung trong cùng đơn hàng (OrderDetail theo order_id)",
             "Gợi ý \"mua kèm\" ở giỏ hàng — rất giá trị cho grocery (mua thịt bò → hành, "
             "tỏi)"),
        ],
    )
    p(d, "Đã cân nhắc nhưng chưa triển khai (ghi nhận gap, không phải ưu tiên hiện tại): "
         "dwell-time trên trang chi tiết (độ chính xác thấp), brand affinity (Product chưa "
         "có field brand riêng), cá nhân hoá theo địa lý (User.address là free-text không "
         "cấu trúc).")

    h2(d, "5. Roadmap theo mốc dữ liệu")
    p(d, "Hệ thống không build 1 lần xong — mỗi phase chỉ bật khi dữ liệu thật đã đạt "
         "ngưỡng tối thiểu, tránh ML thật cho kết quả nhiễu/vô nghĩa trên dữ liệu quá thưa.")
    data_table(
        d,
        ["Phase", "Ngưỡng dữ liệu tối thiểu để bật", "Bổ sung chính"],
        [
            ("0", "Không cần — chỉ instrumentation + rule-based",
             "Bảng tracking mới, sessionId + merge-on-login, fix tạm \"sản phẩm tương tự\""),
            ("1", "2-4 tuần instrumentation ổn định (~500-1000 view, ~50-100 search)",
             "Python service, content-based + co-purchase + popularity, MMR diversity"),
            ("2", "~200-300 đơn PAID, ~50 user có ≥2 đơn, 4-8 tuần view/impression liên tục",
             "Collaborative filtering (ALS), category-fatigue decay, time-decay profile"),
            ("3", "Ổn định 4-8 tuần, ≥2000-3000 impression/tuần từ nguồn AI thật",
             "LinUCB bandit explore-exploit, conversion attribution, A/B testing, dashboard"),
        ],
    )
    p(d, "Trạng thái hiện tại: cả 4 phase đã code xong và verify E2E — xem Phần IV cho chi "
         "tiết lịch sử từng phase, Phần VI cho tình trạng dữ liệu vận hành hiện tại.")

    # ============================================================ PHẦN II
    h1(d, "PHẦN II — CÁC THUẬT TOÁN & KỸ THUẬT")
    p(d, "Mỗi thuật toán được giải thích một lần duy nhất ở đây — tránh lặp lại rải rác "
         "theo từng phase như bản trước. Phần IV (lịch sử triển khai) sẽ chỉ tham chiếu "
         "ngược lại mục tương ứng.")

    h2(d, "1. TF-IDF + Cosine Similarity — gợi ý theo nội dung")
    p(d, "Biến tên, mô tả, danh mục và phân khúc giá của mỗi sản phẩm thành một \"dấu vân "
         "tay\" số học thể hiện sản phẩm đó \"nói về điều gì\" (TF-IDF). Cosine similarity "
         "đo độ giống nhau giữa 2 dấu vân tay — sản phẩm càng giống nhau về nội dung thì "
         "điểm càng cao. Dùng cho \"Sản phẩm tương tự\" và tìm sản phẩm giống các sản phẩm "
         "khách vừa xem. Ưu điểm lớn nhất: giải được cold-start cho sản phẩm mới (không cần "
         "chờ ai mua/xem nó).")
    code_box(d, """
vectorizer = TfidfVectorizer(
    lowercase=True, ngram_range=(1, 2), max_features=20000, sublinear_tf=True
)
matrix = vectorizer.fit_transform(mo_ta_san_pham)  # mỗi hàng = 1 sản phẩm
""")

    h2(d, "2. Association rule mining — gợi ý mua kèm")
    p(d, "Bài toán \"market basket\": quét lịch sử đơn hàng để tìm sản phẩm nào thường được "
         "mua chung (mua thịt bò → thường kèm hành, tỏi), sinh ra luật dạng \"khách mua A "
         "→ gợi ý B\". Apriori là thuật toán kinh điển cho bài toán này nhưng quét lại "
         "toàn bộ dữ liệu nhiều lần nên chậm với tập dữ liệu lớn; FP-Growth đạt cùng kết "
         "quả bằng cách nén giao dịch vào một cấu trúc cây (FP-tree) rồi chỉ quét 2 lần — "
         "đây là thuật toán được chọn dùng thật trong hệ thống.")
    p(d, "Cần đủ số lượng đơn hàng mới cho kết quả có ý nghĩa (tối thiểu 30 đơn trong cửa "
         "sổ 180 ngày); nếu chưa đủ, nguồn này tự \"im lặng\" — không đưa ra luật sai, "
         "không báo lỗi.")
    code_box(d, """
if so_luong_don_hang < 30:
    return {}   # "im lặng" — chưa đủ dữ liệu, không suy diễn sai

tap_pho_bien = fpgrowth(du_lieu_don_hang, min_support=nguong_toi_thieu, max_len=2)
luat = association_rules(tap_pho_bien, metric="confidence", min_threshold=0.1)
""")

    h2(d, "3. Điểm phổ biến (Popularity/Trending)")
    p(d, "Thống kê có trọng số giữa số lượng đã bán và xu hướng mua 7/30 ngày gần nhất. "
         "Đây là nguồn \"an toàn\" duy nhất luôn có kết quả, kể cả với khách vãng lai hoặc "
         "sản phẩm hoàn toàn mới — làm lưới an toàn cho mọi nguồn khác.")
    code_box(d, """
# Điểm phổ biến = 50% mức đã bán + 35% xu hướng 7 ngày + 15% xu hướng 30 ngày
score = (
    0.5 * chuan_hoa(da_ban)
    + 0.35 * chuan_hoa(xu_huong_7_ngay)
    + 0.15 * chuan_hoa(xu_huong_30_ngay)
)
""")

    h2(d, "4. MMR — đa dạng hoá trong 1 lần hiển thị")
    p(d, "Maximal Marginal Relevance: tái xếp hạng bằng cách chọn tuần tự từng sản phẩm sao "
         "cho vừa có điểm liên quan cao vừa không quá giống các sản phẩm đã chọn trước đó "
         "trong cùng danh sách — tránh tình trạng 12 sản phẩm gợi ý đều là 12 loại gạo "
         "giống nhau. Hệ số λ quyết định mức đánh đổi giữa \"liên quan\" và \"đa dạng\"; λ "
         "tự hạ (ép đa dạng mạnh hơn) cho khách có lịch sử xem hẹp, dựa trên category-"
         "entropy — một đại lượng đo \"độ đa dạng\" trong lý thuyết thông tin. Bật từ Phase "
         "1, không đợi có đủ hành vi cá nhân vì chỉ cần category/nội dung là tính được.")
    code_box(d, """
# Ở lượt chọn đầu tiên, chưa có gì được chọn nên chỉ tính điểm liên quan thuần tuý
# -> sản phẩm điểm cao nhất luôn được chọn đầu tiên
gia_tri = lambda_da_dang * diem_lien_quan - (1 - lambda_da_dang) * do_giong_voi_da_chon
""")

    h2(d, "5. DPP — phương án thay thế MMR (chưa dùng)")
    p(d, "Determinantal Point Process: một kỹ thuật đa dạng hoá xác suất khác, chọn cả "
         "tập hợp sản phẩm cùng lúc (thay vì chọn tuần tự từng sản phẩm như MMR) sao cho "
         "\"thể tích\" mà tập đó trải ra trong không gian đặc trưng là lớn nhất — về lý "
         "thuyết cho kết quả đa dạng mượt hơn MMR vì tối ưu toàn cục thay vì tham lam từng "
         "bước, nhưng tính toán tốn kém hơn (cần phân rã ma trận). Roadmap ghi nhận đây là "
         "phương án cân nhắc cho tương lai, chỉ chuyển sang dùng nếu số liệu đo lường thực "
         "tế (xem Phần III mục 2) cho thấy MMR chưa đủ đa dạng — hiện tại chưa có bằng "
         "chứng đó nên chưa triển khai.")

    h2(d, "6. ALS — Collaborative Filtering")
    p(d, "Trả lời câu hỏi khác hẳn TF-IDF: \"những khách hàng có gu mua sắm giống nhau "
         "thường thích sản phẩm gì\" — bất kể 2 sản phẩm đó có mô tả giống nhau hay không. "
         "Mỗi khách hàng và mỗi sản phẩm được biểu diễn bằng một vector số (\"gu sở "
         "thích\"/\"đặc điểm được ưa chuộng\"). Alternating Least Squares học 2 vector này "
         "bằng cách chỉnh xen kẽ — cố định sản phẩm để chỉnh khách hàng cho khớp, rồi cố "
         "định khách hàng để chỉnh sản phẩm, lặp lại nhiều vòng — sao cho tích 2 vector "
         "càng gần với việc khách có thực sự tương tác với sản phẩm đó hay không.")
    p(d, "Chỉ bật khi đủ dữ liệu (guard ≥50 khách có hành vi, ≥20 sản phẩm liên quan) vì "
         "ma trận quá thưa cho kết quả nhiễu. Thư viện `implicit` chuẩn không build được "
         "wheel trên Windows/Python 3.12/numpy2 — hệ thống dùng một bản triển khai ALS "
         "tương đương tự viết (thuật toán Hu-Koren-Volinsky) làm đường chính thức.")
    picture(d, IMG_ALS, "Sơ đồ 2 — Cách ALS học \"gu sở thích\" từ hành vi mua/xem/giỏ hàng")
    code_box(d, """
if so_khach_du_du_lieu < 50 or so_san_pham_lien_quan < 20:
    return None   # "im lặng" -- chưa đủ dữ liệu, chưa bật hành vi cộng đồng

for vong_lap in range(so_vong_lap):
    vector_khach = giai_theo(vector_san_pham, diem_hanh_vi)
    vector_san_pham = giai_theo(vector_khach, diem_hanh_vi)

# Sản phẩm "tương tự theo hành vi" và gợi ý cá nhân hoá suy ra từ 2 vector đã học:
do_giong_hanh_vi(sp_a, sp_b) = cosine(vector_san_pham[sp_a], vector_san_pham[sp_b])
diem_ca_nhan_hoa[san_pham] = vector_khach . vector_san_pham[san_pham]
""")

    h2(d, "7. Category-fatigue decay — chống nhàm chán xuyên nhiều lần hiển thị")
    p(d, "Khác với MMR (chỉ so sánh trong 1 lần gợi ý), cơ chế này nhớ lại nhiều lần gợi ý "
         "trước đó của cùng 1 khách (đọc từ `RecommendationImpression`, cửa sổ 14 ngày). "
         "Nếu 1 danh mục đã được cho xem nhiều lần gần đây mà khách không bấm vào, điểm "
         "các sản phẩm cùng danh mục đó bị hạ thấp ở lần gợi ý tiếp theo — phá vỡ vòng lặp "
         "\"cứ hiện hoài 1 vài loại sản phẩm quen thuộc\".")
    code_box(d, """
diem_nham_chan[danh_muc] += exp(-so_ngay_da_qua * ln2 / 7)  # mỗi lần đã hiển thị mà không bấm

diem_sau_phat = diem_truoc_phat / (1 + 0.15 * diem_nham_chan[danh_muc_cua_san_pham])
""")

    h2(d, "8. Time-decay preference vector & chu kỳ mua lại")
    p(d, "Hồ sơ sở thích trang chủ tính điểm ưu tiên cho từng sản phẩm ứng viên dựa trên "
         "lịch sử tương tác gần đây, có tính \"độ mới\": tương tác càng gần đây càng ảnh "
         "hưởng nhiều, tương tác cũ dần bị \"quên\" theo hàm mũ. Mỗi loại hành vi có tốc "
         "độ quên riêng — mua thật nhớ lâu nhất (30 ngày), bỏ giỏ (14 ngày), xem quên "
         "nhanh nhất (7 ngày).")
    code_box(d, """
# Mua hàng quan trọng nhất và được nhớ lâu nhất (30 ngày);
# xem chỉ là tín hiệu yếu và bị quên nhanh (7 ngày)
diem_hanh_vi = (
    5.0 * exp(-so_ngay_da_qua_khi_mua * ln2 / 30)
    + 2.0 * exp(-so_ngay_da_qua_khi_bo_gio * ln2 / 14)
    + 1.0 * exp(-so_ngay_da_qua_khi_xem * ln2 / 7)
)
""")
    p(d, "Riêng tín hiệu \"mua thật\" cần thêm 1 điều chỉnh đặc thù grocery: khách vừa mua "
         "xong một sản phẩm thường CHƯA cần mua lại ngay (hiệu ứng no/satiation) — công "
         "thức trên nếu dùng nguyên bản sẽ boost mạnh nhất ngay lúc vừa mua, ngược hoàn "
         "toàn với thực tế. Hệ thống suy ra \"chu kỳ mua lại\" ước tính cho từng sản phẩm "
         "(median khoảng cách giữa các lần mua lặp lại của cùng khách, fallback cấp danh "
         "mục rồi hằng số mặc định 14 ngày), rồi nhân thêm hệ số \"sẵn sàng mua lại\" — "
         "thấp ngay sau khi mua, tăng dần tới 1.0 khi gần/qua chu kỳ ước tính.")
    code_box(d, """
san_sang_mua_lai = min(so_ngay_da_qua / chu_ky_mua_lai_uoc_tinh, 1.0)
trong_so_mua_that = 5.0 * exp(-so_ngay_da_qua * ln2 / 30) * san_sang_mua_lai
""")

    h2(d, "9. LinUCB bandit — khám phá-khai thác (explore-exploit)")
    p(d, "Mọi thuật toán ở mục 1-8 đều dựa trên những gì khách ĐÃ làm — điểm mù là hệ "
         "thống không bao giờ biết khách có thích một danh mục hoàn toàn mới hay không, vì "
         "danh mục đó chưa từng được hiển thị đủ để có dữ liệu (\"bong bóng lọc\" tự củng "
         "cố). LinUCB chủ động dành 1-2 vị trí trong mỗi danh sách gợi ý (vị trí thứ 3 và "
         "8, tránh ô cuối vì khách ít nhìn tới) cho một sản phẩm thuộc danh mục ngoài vùng "
         "quen thuộc.")
    p(d, "Mỗi danh mục là một \"cánh tay\" (arm) của bandit. Thuật toán cân nhắc giữa "
         "\"khai thác\" (danh mục từng được đón nhận tốt trong bối cảnh tương tự — điểm kỳ "
         "vọng cao) và \"khám phá\" (danh mục còn ít dữ liệu nên độ bất định cao — đáng thử "
         "để học thêm, cộng thêm điểm \"tò mò\"). Bối cảnh gồm giờ trong ngày, thứ trong "
         "tuần, khách mới hay quen, độ rộng sở thích, và mức độ chán hiện tại với chính "
         "danh mục đó — nhờ chiều bối cảnh cuối này, LinUCB tự học được quy luật \"danh "
         "mục đang bị chán thì thưởng thấp\" mà không cần viết luật riêng. Thompson "
         "Sampling là phương án thay thế được cân nhắc ban đầu — thay vì tính một khoảng "
         "tin cậy tường minh như LinUCB, nó lấy mẫu ngẫu nhiên từ phân phối xác suất của "
         "từng cánh tay rồi chọn cánh tay có mẫu cao nhất; LinUCB được chọn vì công thức "
         "tường minh dễ kiểm chứng/debug hơn khi mới triển khai.")
    picture(d, IMG_LINUCB, "Sơ đồ 3 — Cách LinUCB quyết định thử danh mục nào cho khe khám phá")
    p(d, "Vòng lặp phần thưởng học gần như tức thời (không chờ retrain hằng đêm): một tác "
         "vụ nền đọc bảng lượt hiển thị mỗi 2 phút và quy đổi phản hồi của khách:")
    bullet(d, "Bấm vào sản phẩm khám phá: +0.2 — tín hiệu quan tâm.")
    bullet(d, "Đặt hàng sản phẩm đó (cửa sổ 7 ngày): +1.0 — tín hiệu mạnh nhất.")
    bullet(d, "Hiển thị hơn 24 giờ mà không được bấm: −0.05 — cố ý nhẹ hơn nhiều so với "
              "mức −0.5 dự kiến ban đầu, vì giao diện không có nút \"bỏ qua\" rõ ràng và "
              "mức phạt lớn sẽ nhấn chìm tín hiệu bấm khi tỉ lệ bấm nền chỉ ~4%.")
    bullet(d, "Chưa từng hiện ra màn hình khách (chưa cuộn tới) sau 6 giờ: bỏ qua, không "
              "thưởng không phạt.")
    p(d, "Trạng thái học (ma trận từng cánh tay + nhật ký quyết định chờ phần thưởng) lưu "
         "trong 1 file SQLite cục bộ của Python — khởi động lại không mất kiến thức đã "
         "học, và vẫn giữ nguyên nguyên tắc \"Python chỉ đọc MySQL\".")
    picture(d, IMG_BANDIT_PIPE, "Sơ đồ 4 — Pipeline gợi ý với khe khám phá và vòng lặp phần thưởng")

    h2(d, "10. CRC32 hashing — chia cohort A/B")
    p(d, "Để so sánh \"bật bandit có thực sự tốt hơn không\", khách được chia cố định "
         "thành 2 nhóm bằng mã băm CRC32 của định danh (user_id khi đăng nhập, session_id "
         "khi vãng lai) modulo 100 so với 1 ngưỡng phần trăm cấu hình được. Cách băm này "
         "tính lại được ngay trong câu lệnh SQL khi phân tích (MySQL có hàm `CRC32()` cùng "
         "chuẩn với `zlib.crc32` phía Python) nên không cần thêm cột \"variant\" nào vào "
         "cơ sở dữ liệu — hạn chế duy nhất là khách vãng lai đăng nhập xong bị tính lại "
         "theo hash mới (~50% khả năng đổi nhóm), chấp nhận được vì là nhiễu nhỏ.")

    # ============================================================ PHẦN III
    h1(d, "PHẦN III — PIPELINE HIỆN TẠI & VÒNG LẶP ĐO LƯỜNG")
    p(d, "Phần này mô tả bức tranh cuối cùng sau khi cả 4 phase đã gộp lại — một yêu cầu "
         "gợi ý hiện tại đi qua tuần tự các bước sau (không lặp lại theo từng phase nữa, "
         "xem Phần IV nếu cần biết thứ tự lịch sử từng bước được thêm vào khi nào).")

    h2(d, "1. Luồng xử lý 1 yêu cầu gợi ý")
    num(d, "Lấy song song 4 danh sách ứng viên: nội dung (TF-IDF), mua kèm (FP-Growth), "
           "phổ biến, và hành vi cộng đồng (ALS) — mỗi nguồn vài chục sản phẩm kèm điểm "
           "riêng, đã lọc `active=true AND quantity>0` ngay từ query gốc.")
    num(d, "Trộn 4 danh sách theo tỉ trọng cấu hình riêng cho từng vị trí hiển thị (trang "
           "chủ/trang sản phẩm/giỏ hàng), dedupe, ra một danh sách chung ~80 sản phẩm.")
    num(d, "Hạ điểm các sản phẩm thuộc danh mục khách đã \"xem hoài mà không bấm\" gần đây "
           "(category-fatigue decay).")
    num(d, "Đa dạng hoá trong-1-lần-hiển-thị (MMR), chọn ra 6-12 sản phẩm cuối cùng.")
    num(d, "Nếu khách thuộc cohort bandit-on: chèn 1-2 khe khám phá vào vị trí 3/8 (thay "
           "item điểm thấp nhất ở các vị trí đó), ghi quyết định vào state bandit.")
    num(d, "Trả về danh sách cuối cùng (mã sản phẩm + điểm + nguồn) cho Java — Java hydrate "
           "lại giá/khuyến mãi/tồn kho mới nhất từ DB trước khi trả về khách (lớp lọc tồn "
           "kho thứ 2, luôn tươi hơn Python vì Python chỉ tính lại theo chu kỳ).")
    p(d, "Nếu bất kỳ bước nào ở Python lỗi/timeout/rỗng, Java chuyển ngay sang gợi ý "
         "rule-based (cùng danh mục/bán chạy) — khách không bao giờ thấy trang trống hay "
         "lỗi.")

    h2(d, "2. Vòng lặp đo lường continual learning")
    p(d, "Để biết mô hình có thực sự cải thiện theo thời gian (không chỉ \"cảm giác\"), "
         "hệ thống đo 6 chỉ số, hiển thị trong trang quản trị \"Gợi ý AI\":")
    bullet(d, "CTR/CVR theo vị trí hiển thị × nguồn thuật toán — so sánh hiệu quả từng "
              "nguồn (content-based/co-purchase/popularity/CF/bandit-explore).")
    bullet(d, "Intra-list diversity = 1 − cosine similarity trung bình theo cặp trong cùng "
              "1 lần hiển thị.")
    bullet(d, "Category entropy của cái được hiển thị vs cái được click theo tuần — phải "
              "tăng dần nếu đa dạng hoá hiệu quả.")
    bullet(d, "Catalog coverage = % sản phẩm active có ≥1 lượt hiển thị/30 ngày — chống "
              "việc mô hình chỉ lặp vài chục sản phẩm quen.")
    bullet(d, "Repeat-impression-no-click rate theo danh mục — phải giảm dần sau khi "
              "fatigue-decay phát huy tác dụng.")
    bullet(d, "So sánh cohort A/B bandit-on vs bandit-off (CTR/CVR + mức chênh lệch).")
    p(d, "Số liệu do Python tính trực tiếp từ cơ sở dữ liệu (cache 5 phút), Java làm cầu "
         "nối và khoá quyền quản trị viên.")

    data_table(
        d,
        ["Thành phần", "Tần suất học lại", "Ghi chú"],
        [
            ("Content-based similarity", "Nightly", "Catalog đổi hàng ngày qua admin CRUD"),
            ("Co-purchase rules", "Nightly → Weekly khi volume lớn", ""),
            ("Collaborative filtering (ALS)", "Nightly khi dữ liệu ít → carry-forward tới "
             "7 ngày khi dữ liệu lớn", "Atomic swap, không có trạng thái nửa cũ nửa mới"),
            ("Bandit params", "Không batch — poll bảng impression mỗi 1-5 phút",
             "Không cần Kafka/message broker ở quy mô hiện tại"),
        ],
    )

    # ============================================================ PHẦN IV
    h1(d, "PHẦN IV — LỊCH SỬ TRIỂN KHAI THEO PHASE")
    p(d, "Tóm tắt những gì mỗi phase đã thêm vào, theo đúng thứ tự thực hiện. Thuật toán "
         "đã giải thích chi tiết ở Phần II — mục này chỉ nhắc lại NGẮN GỌN phần nào mới, "
         "không lặp lại công thức.")

    h2(d, "Phase 0 — Instrumentation + rule-based")
    p(d, "Chưa cần Python service. Xây 4 bảng theo dõi hành vi mới (`ProductView`, "
         "`SearchLog`, `RecommendationImpression`, `CartEvent`), sessionId ẩn danh cho "
         "guest + merge khi đăng nhập, và sửa tạm \"Sản phẩm tương tự\" bằng rule Java "
         "thuần (cùng danh mục, sort theo bán chạy/đánh giá) để hết gãy trải nghiệm ngay "
         "trong lúc chờ đủ dữ liệu cho Phase 1.")

    h2(d, "Phase 1 — Content-based + Co-purchase + Popularity")
    p(d, "Dựng Python service, bật content-based + popularity ngay (không phụ thuộc hành "
         "vi cá nhân), co-purchase chạy song song im lặng tới khi đủ dữ liệu. Bật MMR "
         "ngay từ đầu — không đợi có CF — để tránh mô hình \"tập thói quen xấu\" sớm. "
         "Thêm 3 vị trí hiển thị (Home/PDP/Cart) dùng chung 1 component `RecommendationRail` "
         "tự ghi impression (IntersectionObserver) và click-through.")
    picture(d, IMG_PIPE1, "Sơ đồ 5 — Quy trình gợi ý 4 bước của Phase 1")
    p(d, "Lỗi phát hiện sau triển khai: DTO trả về gợi ý thiếu field tồn kho, khiến nút "
         "\"thêm nhanh vào giỏ\" luôn báo sai hết hàng VÀ sản phẩm hết hàng vẫn lọt vào "
         "danh sách gợi ý — đây là lần đầu tiên của một lớp bug lặp lại nhiều lần, xem "
         "Phần V mục 1 để biết đầy đủ 3 lần xảy ra và bài học chung.")

    h2(d, "Phase 2 — Collaborative Filtering + fatigue decay")
    p(d, "Bật ALS thật (ngưỡng đạt: 58 khách hàng, 228 sản phẩm — học xong trong 0,28 "
         "giây), category-fatigue decay, và mở rộng hồ sơ sở thích trang chủ sang cả "
         "bỏ giỏ/mua thật (trước đó Phase 1 chỉ tính từ lịch sử xem). Java/frontend cập "
         "nhật gửi kèm `user_id`/`session_id` ở cả 3 vị trí hiển thị (trước đó chỉ trang "
         "chủ gửi).")
    picture(d, IMG_PIPE2, "Sơ đồ 6 — Quy trình gợi ý 5 bước sau khi thêm CF + chống nhàm chán theo "
                          "lịch sử (Phase 2)")
    p(d, "2 lỗi phát hiện sau triển khai (qua quan sát UI thật, không phải kiểm thử tự "
         "động): (1) cùng lớp bug thiếu lọc tồn kho như Phase 1, lần này ở nguồn CF — xem "
         "Phần V mục 1; (2) tín hiệu \"vừa mua\" được coi là mạnh nhất ngay lập tức thay vì "
         "yếu dần rồi tăng lại theo chu kỳ mua lại — đã dẫn tới việc xây module "
         "\"repurchase readiness\" mô tả ở Phần II mục 8.")
    p(d, "Kiểm thử đã thực hiện: xác nhận đúng ngưỡng dữ liệu; giả lập \"chưa đủ dữ liệu\" "
         "→ CF tự tắt hoàn toàn, không lỗi; giả lập chống nhàm chán (17 lần hiển thị không "
         "click 1 danh mục → tỉ lệ giảm từ 8/8 xuống 4/8, đánh dấu đã click → hồi 8/8); tắt "
         "hẳn Python service → cả 3 vị trí vẫn trả kết quả trong <200ms qua fallback.")

    h2(d, "Phase 3 — LinUCB bandit + A/B testing + dashboard")
    p(d, "Bổ sung khe khám phá LinUCB (Phần II mục 9), ghi nhận conversion (khoảng trống "
         "tồn tại từ Phase 0: cột `converted` có sẵn trong schema nhưng chưa từng được "
         "ghi), A/B cohort qua CRC32 (Phần II mục 10), và dashboard đo lường (Phần III "
         "mục 2).")
    p(d, "2 việc nền tảng phải làm trước khi bandit hoạt động đúng: (1) chuyển việc sinh "
         "`request_id` từ Java (sau khi nhận response) sang Python (trước khi tính toán) "
         "để Python join được quyết định bandit với impression sau này; (2) thêm nguồn "
         "thuật toán theo từng item (`itemSources`) thay vì chỉ theo cả slate, vì khe khám "
         "phá cần nhãn `BANDIT_EXPLORE` riêng khác với phần còn lại của danh sách.")
    p(d, "Kiểm thử đã thực hiện qua HTTP thật xuyên suốt 3 tầng: khe khám phá đúng vị trí "
         "3/8 chỉ với cohort bật; đủ 4 nhánh vòng lặp phần thưởng (bấm/mua/im lặng quá "
         "hạn/chưa hiện ra màn hình); đặt đơn COD thật xác nhận conversion đúng sản phẩm "
         "đúng đơn; vòng lặp khép kín thật (bấm → dưới 2 phút sau nhật ký ghi nhận phần "
         "thưởng); phân quyền trang số liệu (guest 401, user thường bị chặn, admin xem đủ "
         "6 khối).")
    p(d, "Công tắc an toàn không cần sửa code: `BANDIT_ENABLED=false` hoặc "
         "`AB_BANDIT_PCT=0` trong `.env` của Python service tắt toàn bộ cơ chế khám phá, "
         "hệ thống quay về đúng hành vi Phase 2.")

    # ============================================================ PHẦN V
    h1(d, "PHẦN V — SỰ CỐ ĐÃ PHÁT HIỆN & BÀI HỌC")
    p(d, "Thay vì kể lại từng lần phát hiện lỗi rải rác theo phase, phần này gộp các lỗi "
         "cùng một lớp nguyên nhân lại với nhau để bài học được nêu rõ một lần, áp dụng "
         "được cho cả những nguồn dữ liệu tương tự thêm vào sau này.")

    h3(d, "1. Lớp bug lặp lại: thiếu lọc active/quantity>0 ngay từ nguồn")
    p(d, "Xảy ra 3 lần độc lập trong suốt quá trình xây dựng — mỗi lần một nguồn ứng viên "
         "MỚI được thêm vào lại quên áp dụng điều kiện lọc tồn kho đã biết từ trước:")
    bullet(d, "Phase 1: cả 3 nguồn gốc (content-based/co-purchase/popularity) cùng DTO trả "
              "về cho giao diện đều thiếu field tồn kho — sản phẩm hết hàng vẫn được gợi ý, "
              "nút thêm giỏ hàng báo sai.")
    bullet(d, "Phase 2: nguồn collaborative filtering (CF) mới thêm vào lấy sản phẩm thẳng "
              "từ lịch sử mua/xem/giỏ mà không lọc lại tồn kho như 3 nguồn kia đã làm — "
              "khiến danh sách gợi ý ngắn hơn số lượng yêu cầu (12 còn 10).")
    bullet(d, "Phát hiện qua code review độc lập (không phải qua kiểm thử theo tính năng): "
              "nguồn co-purchase (đã tồn tại từ Phase 1) chưa từng lọc tồn kho khi tính "
              "luật kết hợp — luật chứa sản phẩm hết hàng vẫn tồn tại trong bộ nhớ, lãng "
              "phí chỗ trong danh sách dù bị Java lọc bỏ ở bước cuối. Đã sửa: thêm "
              "`JOIN products ... WHERE active=1 AND quantity>0` vào câu truy vấn xây "
              "giỏ hàng giao dịch.")
    p(d, "Bài học: điều kiện lọc tồn kho phải là một bước bắt buộc trong checklist khi "
         "thêm BẤT KỲ nguồn ứng viên mới nào, không chỉ tin vào lớp lọc cuối cùng ở Java "
         "(lớp đó chỉ đảm bảo khách không thấy sản phẩm hết hàng, không đảm bảo danh sách "
         "đủ số lượng hay không lãng phí chỗ trong pool ứng viên).")

    h3(d, "2. Lớp bug lặp lại: nhầm giờ địa phương với UTC")
    p(d, "`order_time` trong cơ sở dữ liệu lưu dạng UTC naive (kiểu `Instant` phía Java), "
         "trong khi `datetime.now()` mặc định của Python trả về giờ hệ thống (giờ địa "
         "phương, UTC+7 tại Việt Nam) nếu không truyền `timezone.utc`. Lệch 7 giờ này từng "
         "xuất hiện ở module tính luật mua kèm (`co_purchase.py`, phát hiện và sửa trong "
         "quá trình xây Phase 2), và phát hiện lại lần thứ hai ở module tính điểm phổ biến "
         "(`popularity.py`) qua code review độc lập sau khi toàn bộ 4 phase đã \"xong\" — "
         "cùng lớp lỗi nhưng không được rà soát chéo sang các file khác dùng thời gian "
         "tương tự khi sửa lần đầu.")
    p(d, "Bài học: mọi hàm mới tính khoảng thời gian từ `order_time`/`shown_at` phải dùng "
         "`datetime.now(timezone.utc).replace(tzinfo=None)`, không dùng `datetime.now()` "
         "trần — nên có 1 lượt rà soát toàn bộ file dùng `datetime.now()` mỗi khi phát "
         "hiện lại lớp lỗi này, thay vì chỉ sửa đúng chỗ vừa tìm thấy.")

    h3(d, "3. Bài học nghiệp vụ: \"vừa mua\" không phải tín hiệu \"muốn mua lại ngay\"")
    p(d, "Không phải bug kỹ thuật mà là một giả định nghiệp vụ sai ban đầu: hành vi mua "
         "sắm grocery có tính chu kỳ, khách vừa mua một sản phẩm thường CHƯA cần mua lại "
         "ngay. Vấn đề này đã được ghi nhận từ nghiên cứu tín hiệu ban đầu (Phần I mục 4) "
         "nhưng chưa được xây tới tận Phase 2, dẫn tới trang chủ ưu tiên gợi ý lại đúng "
         "sản phẩm khách vừa mua xong. Cách khắc phục (\"repurchase readiness\") mô tả ở "
         "Phần II mục 8.")

    h3(d, "4. Bảo mật: định danh phải luôn lấy từ JWT, không tin tham số client")
    p(d, "Phát hiện qua code review độc lập sau khi Phase 3 đã \"xong\": endpoint đặt hàng "
         "tạo `Order` đúng từ JWT, nhưng dòng gọi ghi nhận conversion ngay sau đó lại dùng "
         "thẳng tham số `userId` do client tự gửi trong request thay vì định danh JWT thật "
         "— một khách đăng nhập có thể gửi kèm `userId` của người khác trong request đặt "
         "hàng để đơn của mình vô tình (hoặc cố ý) đánh dấu \"chuyển đổi\" lên lượt hiển "
         "thị gợi ý của người đó, làm sai lệch phần thưởng bandit và số liệu CVR trên "
         "dashboard. Đơn hàng chính nó không bị ảnh hưởng (vẫn luôn tạo đúng theo JWT) — "
         "chỉ riêng bước ghi nhận conversion dùng nhầm biến. Đã sửa: dùng lại "
         "`order.getUser().getId()` (đã derive từ JWT khi tạo đơn) thay vì tham số request.")
    p(d, "Bài học: nguyên tắc \"định danh chỉ lấy từ JWT\" (Phần I mục 3) phải áp dụng cho "
         "MỌI lời gọi dùng định danh người dùng trong cùng 1 luồng xử lý, không chỉ lời "
         "gọi đầu tiên (tạo đơn) — dễ sót ở các bước phụ được thêm sau (ở đây là "
         "attribution, thêm vào Phase 3, dùng lại tham số controller đã có sẵn từ trước "
         "thay vì lấy lại đúng nguồn).")

    h3(d, "5. Hiệu năng: N+1 query khi ghi nhận hàng loạt lượt hiển thị")
    p(d, "Phát hiện qua code review độc lập: mỗi lần một danh sách gợi ý lọt vào viewport, "
         "frontend gửi 1 batch tối đa 50 lượt hiển thị; endpoint ghi nhận cũ xử lý từng "
         "item bằng 1 câu kiểm tra tồn tại + 1 câu ghi riêng lẻ (tối đa 100 câu truy vấn "
         "cho 1 request). Đã sửa: gộp thành 1 câu kiểm tra tồn tại theo lô (`findAllById`) "
         "và 1 câu ghi theo lô (`saveAll`).")

    h3(d, "6. Đã biết, chấp nhận không sửa (không phải bug, hành vi mặc định của hệ thống)")
    bullet(d, "`cart-suggestions?productIds=abc` (chữ thay vì số) → lỗi 500 type-mismatch "
              "— hành vi mặc định toàn hệ thống với tham số sai kiểu, frontend không bao "
              "giờ gửi được giá trị này.")
    bullet(d, "`products/similar/{id không tồn tại}` → trả về danh sách phổ biến thay vì "
              "rỗng — fail-safe hợp lý, trang chi tiết sản phẩm đã tự trả 404 trước đó rồi.")
    bullet(d, "LinUCB dùng state in-memory dạng singleton — chỉ an toàn khi chạy đúng 1 "
              "worker process (`uvicorn` mặc định). Chưa cần scale ngang ở quy mô hiện tại "
              "nên chưa refactor; đã ghi rõ ràng buộc này trong README vận hành.")

    # ============================================================ PHẦN VI
    h1(d, "PHẦN VI — KIỂM CHỨNG & VẬN HÀNH HIỆN TẠI")

    h2(d, "1. Tình trạng dữ liệu")
    p(d, "Cả Phase 2 và Phase 3 hiện đang chạy trên một phần dữ liệu MẪU được bơm sẵn "
         "(script + manifest kèm theo) để đạt đủ ngưỡng dữ liệu tối thiểu và kiểm tra khả "
         "năng dự đoán trước khi có đủ dữ liệu thật:")
    bullet(d, "Phase 2: ~40 user + ~104 đơn PAID + event tracking trải 4 tuần, đưa tổng số "
              "lên 284 đơn PAID / 56 user có ≥2 đơn.")
    bullet(d, "Phase 3: 8 tuần × ~2.600-2.800 lượt hiển thị nguồn AI mỗi tuần + 230 lượt "
              "chuyển đổi neo vào đơn hàng thật.")
    p(d, "Cả 2 script seed đều ghi lại manifest JSON (danh sách chính xác các dòng đã "
         "chèn) để xoá sạch bằng script dọn dẹp dùng chung khi không còn cần nữa. Khuyến "
         "nghị: xoá dữ liệu mẫu Phase 3 TRƯỚC khi bắt đầu tin các con số của bandit như số "
         "liệu thật — trong thời gian thử nghiệm, bandit học trực tiếp trên phản hồi "
         "(click/mua) của chính dữ liệu mẫu này.")

    h2(d, "2. Công tắc an toàn (kill-switch), không cần sửa code")
    bullet(d, "`BANDIT_ENABLED=false` hoặc `AB_BANDIT_PCT=0` — tắt toàn bộ khe khám phá, "
              "hệ thống quay về đúng hành vi Phase 2.")
    bullet(d, "CF, co-purchase tự tắt im lặng khi dữ liệu tụt dưới ngưỡng guard (không cần "
              "cấu hình gì) — đúng thiết kế \"im lặng khi thiếu dữ liệu\".")
    bullet(d, "Tắt hẳn `recommendation-service`: mọi endpoint gợi ý tự fallback rule-based "
              "trong Java, trang web không vỡ.")

    h2(d, "3. Endpoint & vận hành")
    p(d, "`user_id` luôn do Java derive từ JWT, không bao giờ nhận từ client. Python chỉ "
         "trả `product_id + score + source`, Java hydrate DTO đầy đủ từ DB (giá/khuyến "
         "mãi/tồn kho luôn mới nhất). Chi tiết endpoint, biến môi trường, cách chạy dev "
         "xem `recommendation-service/README.md`.")


def main() -> None:
    d = Document(str(DOCX))
    body = d.element.body
    sect_pr = body.find(qn("w:sectPr"))
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)

    build(d)

    d.save(str(DOCX))
    print("Đã viết lại", DOCX)


if __name__ == "__main__":
    main()
