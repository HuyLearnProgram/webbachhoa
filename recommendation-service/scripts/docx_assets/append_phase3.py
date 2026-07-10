# -*- coding: utf-8 -*-
r"""Bổ sung "PHẦN IV — QUY TRÌNH XÂY DỰNG MÔ HÌNH (PHASE 3)" vào AI_Recommendation_Roadmap.docx.

Tái dùng đúng style các Phần II/III có sẵn trong tài liệu (Heading 1/2/3, List Bullet,
đoạn code Consolas 10pt, Caption cho sơ đồ) + 2 sơ đồ matplotlib mới:
  - Sơ đồ 5: pipeline gợi ý có thêm khe "khám phá" của bandit + vòng lặp phần thưởng.
  - Sơ đồ 6: cách LinUCB quyết định chọn danh mục nào để thử.

Chạy: cd recommendation-service && venv\Scripts\python scripts/docx_assets/append_phase3.py
(Đóng file Word trước khi chạy.)
"""
import sys
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch

from docx import Document
from docx.shared import Pt, Inches

ROOT = Path(__file__).resolve().parents[2]
DOCX = ROOT / "AI_Recommendation_Roadmap.docx"
ASSETS = Path(__file__).resolve().parent

BLUE, GREEN, ORANGE, RED, GRAY = "#2a78d6", "#1baf7a", "#eda100", "#e34948", "#5b5b5b"


# ---------------------------------------------------------------- sơ đồ

def _box(ax, xy, w, h, text, color, fs=10):
    x, y = xy
    ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.06",
                                fc=color, ec="none", alpha=0.14))
    ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.06",
                                fc="none", ec=color, lw=1.6))
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=fs, color="#222",
            wrap=True)


def _arrow(ax, a, b, color=GRAY, style="-|>", lw=1.6, ls="-"):
    ax.add_patch(FancyArrowPatch(a, b, arrowstyle=style, mutation_scale=16,
                                 color=color, lw=lw, linestyle=ls))


def draw_pipeline(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(10.5, 5.6))
    ax.set_xlim(0, 10.5); ax.set_ylim(0, 5.6); ax.axis("off")

    _box(ax, (0.2, 4.3), 2.2, 0.9, "Ứng viên từ 4 nguồn\n(nội dung, mua chung,\nphổ biến, cộng đồng)", BLUE)
    _box(ax, (2.9, 4.3), 1.8, 0.9, "Trộn điểm +\ngiảm nhàm chán", BLUE)
    _box(ax, (5.2, 4.3), 1.6, 0.9, "Đa dạng hoá\n(MMR)", BLUE)
    _box(ax, (7.3, 4.3), 2.9, 0.9, "Chèn 1-2 khe KHÁM PHÁ\n(bandit chọn danh mục lạ,\nchỉ nhóm A/B bật)", GREEN)
    _arrow(ax, (2.4, 4.75), (2.9, 4.75)); _arrow(ax, (4.7, 4.75), (5.2, 4.75))
    _arrow(ax, (6.8, 4.75), (7.3, 4.75))

    _box(ax, (7.3, 2.9), 2.9, 0.8, "Danh sách gợi ý hiển thị\ncho khách (ghi lượt hiển thị)", ORANGE)
    _arrow(ax, (8.75, 4.3), (8.75, 3.7), color=ORANGE)

    _box(ax, (3.9, 2.9), 2.6, 0.8, "Khách phản hồi:\nbấm xem / đặt hàng / bỏ qua", ORANGE)
    _arrow(ax, (7.3, 3.3), (6.5, 3.3), color=ORANGE)

    _box(ax, (3.9, 1.3), 2.6, 0.9, "Vòng lặp phần thưởng\n(đọc bảng lượt hiển thị\nmỗi 2 phút)", RED)
    _arrow(ax, (5.2, 2.9), (5.2, 2.2), color=RED)

    _box(ax, (0.4, 1.3), 2.6, 0.9, "Bandit cập nhật ngay:\nbấm +0.2 · mua +1.0\nbỏ qua −0.05", GREEN)
    _arrow(ax, (3.9, 1.75), (3.0, 1.75), color=RED)
    _arrow(ax, (1.7, 2.2), (7.6, 4.32), color=GREEN, ls=(0, (4, 3)))
    ax.text(3.4, 3.6, "lần gợi ý sau\nthông minh hơn", fontsize=9, color=GREEN, ha="center",
            style="italic")

    fig.tight_layout()
    fig.savefig(path, dpi=160)
    plt.close(fig)


def draw_linucb(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(10.5, 5.2))
    ax.set_xlim(0, 10.5); ax.set_ylim(0, 5.2); ax.axis("off")

    _box(ax, (0.2, 3.6), 2.9, 1.2,
         "Bối cảnh lúc gợi ý:\n• giờ trong ngày, thứ trong tuần\n• khách mới hay quen\n"
         "• độ rộng sở thích\n• mức chán từng danh mục", BLUE, fs=9)

    cats = [("Rau củ", "đã biết rõ:\nhay được bấm", GREEN, 0.55),
            ("Đồ khô", "ít dữ liệu:\nđáng thử", ORANGE, 0.80),
            ("Gia vị", "đã thử nhiều,\nít ai bấm", RED, 0.25)]
    for i, (name, note, color, score) in enumerate(cats):
        x = 4.0 + i * 2.15
        _box(ax, (x, 3.6), 1.85, 1.2, f"{name}\n{note}", color, fs=9)
        ax.text(x + 0.92, 3.30, f"điểm tin cậy + tò mò\n= {score:.2f}", ha="center", fontsize=8.5,
                color=color)
        _arrow(ax, (3.1, 4.2), (x, 4.2), color=GRAY, lw=1.1)

    _box(ax, (3.6, 1.2), 3.4, 1.0,
         "Chọn danh mục điểm cao nhất → lấy 1 sản phẩm\nbán chạy còn hàng của danh mục đó\n"
         "đưa vào khe khám phá", GREEN, fs=9)
    _arrow(ax, (6.15 + 0.92 - 2.15, 3.3), (5.3, 2.2), color=ORANGE)

    ax.text(5.25, 0.55,
            "Điểm = kỳ vọng từ dữ liệu đã học + phần \"tò mò\" (danh mục càng ít dữ liệu phần này càng lớn)\n"
            "→ lúc đầu thử đều mọi danh mục, càng nhiều phản hồi càng dồn về danh mục thực sự hợp gu.",
            ha="center", fontsize=9, color="#333", style="italic")

    fig.tight_layout()
    fig.savefig(path, dpi=160)
    plt.close(fig)


# ---------------------------------------------------------------- docx helpers

def h1(d, t): d.add_paragraph(t, style="Heading 1")
def h2(d, t): d.add_paragraph(t, style="Heading 2")
def h3(d, t): d.add_paragraph(t, style="Heading 3")
def p(d, t): d.add_paragraph(t, style="Normal")
def bullet(d, t): d.add_paragraph(t, style="List Bullet")


def code(d, text):
    para = d.add_paragraph(style="Normal")
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)


def picture(d, img: Path, caption: str):
    d.add_picture(str(img), width=Inches(6.3))
    d.add_paragraph(caption, style="Caption")


# ---------------------------------------------------------------- nội dung

def main() -> None:
    img1 = ASSETS / "phase3_pipeline.png"
    img2 = ASSETS / "phase3_linucb.png"
    draw_pipeline(img1)
    draw_linucb(img2)

    d = Document(str(DOCX))
    if any(pr.text.startswith("PHẦN IV") for pr in d.paragraphs if pr.style.name == "Heading 1"):
        print("PHẦN IV đã tồn tại — không thêm lại.")
        return

    h1(d, "PHẦN IV — QUY TRÌNH XÂY DỰNG MÔ HÌNH (PHASE 3)")
    p(d, "Phase 3 bổ sung mảnh ghép cuối của vòng lặp tự học: cơ chế khám phá–khai thác "
         "(explore–exploit) bằng thuật toán LinUCB, kèm khả năng ghi nhận đơn hàng phát sinh "
         "từ gợi ý, thử nghiệm A/B và bộ số liệu đo lường hiển thị ngay trong trang quản trị. "
         "Nếu Phase 2 giúp hệ thống hiểu khách qua lịch sử, thì Phase 3 cho phép hệ thống chủ "
         "động thử cái mới một cách có kiểm soát và tự điều chỉnh theo phản hồi trong ngày, "
         "thay vì chờ tới lượt huấn luyện lại hằng đêm.")

    h2(d, "1. Vì sao cần khe \"khám phá\" và bandit")
    p(d, "Mọi nguồn gợi ý ở Phase 1–2 đều dựa trên những gì khách ĐÃ làm: đã xem gì, đã mua gì, "
         "người giống mình mua gì. Điểm mù của cách này là hệ thống không bao giờ biết khách có "
         "thích một danh mục hoàn toàn mới hay không, vì danh mục đó chưa từng được hiển thị đủ "
         "nhiều để có dữ liệu — hiệu ứng \"bong bóng lọc\" tự củng cố. Cơ chế chống nhàm chán ở "
         "Phase 2 chỉ giảm bớt cái cũ một cách thụ động; Phase 3 chủ động dành 1-2 vị trí trong "
         "mỗi danh sách gợi ý (vị trí thứ 3 và thứ 8, tránh ô cuối vì khách ít nhìn tới) cho một "
         "sản phẩm thuộc danh mục nằm ngoài vùng quen thuộc.")
    p(d, "Chọn danh mục nào để thử không phải ngẫu nhiên: thuật toán LinUCB cân nhắc giữa "
         "\"khai thác\" (danh mục từng được đón nhận tốt trong bối cảnh tương tự) và \"khám phá\" "
         "(danh mục còn ít dữ liệu nên độ bất định cao — đáng thử để học thêm). Mỗi danh mục là "
         "một \"cánh tay\" của bandit; bối cảnh gồm giờ trong ngày, thứ trong tuần, khách mới hay "
         "quen, độ rộng sở thích và mức độ chán hiện tại với chính danh mục đó. Nhờ chiều bối "
         "cảnh cuối này, LinUCB tự học được quy luật \"danh mục đang bị chán thì thưởng thấp\" — "
         "không cần viết luật riêng.")
    picture(d, img2, "Sơ đồ 6 — Cách LinUCB quyết định thử danh mục nào cho khe khám phá")

    h2(d, "2. Vòng lặp phần thưởng — học ngay trong ngày")
    p(d, "Khác với các nguồn gợi ý huấn luyện lại hằng đêm, bandit cập nhật gần như tức thời: "
         "một tác vụ nền đọc bảng lượt hiển thị mỗi 2 phút (đúng chủ trương trong roadmap — "
         "không thêm message broker ở quy mô hiện tại) và quy đổi phản hồi của khách thành "
         "phần thưởng:")
    bullet(d, "Khách bấm vào sản phẩm khám phá: +0.2 — tín hiệu quan tâm.")
    bullet(d, "Khách đặt hàng sản phẩm đó (trong cửa sổ 7 ngày): +1.0 — tín hiệu mạnh nhất.")
    bullet(d, "Hiển thị hơn 24 giờ mà không được bấm: −0.05 — tín hiệu chán, cố ý để nhẹ hơn "
              "nhiều so với mức −0.5 dự kiến ban đầu vì giao diện không có nút \"bỏ qua\" rõ "
              "ràng, và mức phạt lớn sẽ nhấn chìm tín hiệu bấm (+0.2) khi tỉ lệ bấm nền chỉ "
              "khoảng 4%.")
    bullet(d, "Danh sách chưa từng hiện ra trên màn hình khách (chưa cuộn tới) sau 6 giờ: bỏ "
              "qua, không thưởng không phạt — chưa nhìn thấy thì không phải phản hồi.")
    p(d, "Trạng thái học của bandit (ma trận từng cánh tay + nhật ký quyết định chờ phần thưởng) "
         "được lưu trong một file SQLite cục bộ của dịch vụ Python — dịch vụ khởi động lại không "
         "mất kiến thức đã học, và giữ nguyên nguyên tắc \"Python chỉ đọc MySQL, mọi thao tác ghi "
         "đi qua Java\". Mất file này chỉ khiến bandit học lại từ đầu, không mất dữ liệu nghiệp vụ.")
    picture(d, img1, "Sơ đồ 5 — Pipeline gợi ý Phase 3 với khe khám phá và vòng lặp phần thưởng")

    h2(d, "3. Ghi nhận đơn hàng phát sinh từ gợi ý (conversion)")
    p(d, "Khi rà soát để làm phần thưởng +1.0, phát hiện một khoảng trống tồn tại từ Phase 0: "
         "bảng lượt hiển thị có sẵn cột \"đã chuyển đổi\" nhưng chưa có chỗ nào ghi vào. Phase 3 "
         "bổ sung: ngay khi khách đặt đơn thành công (kể cả COD chưa thanh toán — theo quyết định "
         "nghiệp vụ: đặt đơn đã là tín hiệu ý định mua, đơn huỷ chấp nhận là nhiễu nhỏ), hệ thống "
         "tìm cho mỗi sản phẩm trong đơn đúng MỘT lượt hiển thị gợi ý gần nhất trong 7 ngày "
         "(ưu tiên lượt đã được bấm) và đánh dấu kèm mã đơn hàng. Chỉ đánh dấu một lượt duy nhất "
         "(last-touch) để không thổi phồng tỉ lệ chuyển đổi. Bước này được đặt sau khi tạo đơn và "
         "bọc kín lỗi — dù có trục trặc cũng không bao giờ ảnh hưởng luồng đặt hàng của khách.")

    h2(d, "4. Thử nghiệm A/B và bảng số liệu trong trang quản trị")
    p(d, "Để trả lời câu hỏi \"bật bandit có thực sự tốt hơn không\", khách được chia làm hai "
         "nhóm cố định theo mã băm CRC32 của định danh (khách đăng nhập theo mã người dùng, khách "
         "vãng lai theo mã phiên): 50% có khe khám phá, 50% không. Cách băm này tính lại được "
         "ngay trong SQL khi phân tích nên không cần thêm cột nào vào cơ sở dữ liệu. Trang quản "
         "trị có thêm mục \"Gợi ý AI\" hiển thị:")
    bullet(d, "Tỉ lệ bấm (CTR) và tỉ lệ chuyển đổi (CVR) theo từng vị trí hiển thị × nguồn thuật "
              "toán — nhìn được ngay nguồn khám phá BANDIT_EXPLORE hoạt động ra sao so với các "
              "nguồn còn lại.")
    bullet(d, "So sánh trực tiếp hai nhóm A/B kèm mức chênh lệch.")
    bullet(d, "Độ đa dạng trong từng danh sách (1 − độ tương đồng trung bình giữa các cặp sản "
              "phẩm cùng danh sách).")
    bullet(d, "Độ rộng sở thích (entropy) của những gì được hiển thị so với những gì được bấm, "
              "theo tuần — phải tăng dần nếu đa dạng hoá có hiệu quả.")
    bullet(d, "Độ phủ danh mục hàng: bao nhiêu phần trăm sản phẩm đang bán có ít nhất một lượt "
              "được gợi ý trong 30 ngày — chống việc mô hình chỉ xoay quanh vài chục sản phẩm quen.")
    bullet(d, "Tỉ lệ hiển thị lặp lại mà không được bấm theo danh mục — phải giảm dần nếu cơ chế "
              "chống nhàm chán làm đúng việc.")
    p(d, "Số liệu do dịch vụ Python tính trực tiếp từ cơ sở dữ liệu (có bộ nhớ đệm 5 phút), Java "
         "làm cầu nối và khoá quyền quản trị viên; giao diện dùng lại đúng bộ biểu đồ sẵn có của "
         "trang tổng quan.")

    h2(d, "5. Kiểm chứng trước khi bàn giao")
    p(d, "Toàn bộ Phase 3 được kiểm chứng qua HTTP thật xuyên suốt ba tầng (giả lập đúng hành vi "
         "frontend → Java → MySQL → Python), không chỉ qua unit test:")
    bullet(d, "Khe khám phá xuất hiện đúng vị trí 3 và 8 với nhóm A/B bật, tuyệt đối không xuất "
              "hiện với nhóm tắt; danh sách 6 món chỉ có 1 khe; sản phẩm khám phá xoay vòng chứ "
              "không lặp một món.")
    bullet(d, "Đủ 4 nhánh của vòng lặp phần thưởng: bấm (+0.2, tiếp tục chờ chuyển đổi), đặt hàng "
              "(+1.0, chốt), im lặng quá hạn (−0.05, chốt), chưa hiện ra màn hình (bỏ qua). Khởi "
              "động lại dịch vụ xác nhận trạng thái học còn nguyên.")
    bullet(d, "Đặt một đơn COD thật chứa sản phẩm từng được gợi ý: lượt hiển thị tương ứng được "
              "đánh dấu chuyển đổi kèm đúng mã đơn; sản phẩm không có lượt hiển thị nào thì không "
              "dòng nào bị đánh dấu.")
    bullet(d, "Khép kín vòng lặp thật: bấm vào sản phẩm khám phá trên đường đi thật của frontend "
              "→ chưa đầy 2 phút sau nhật ký dịch vụ ghi nhận phần thưởng và ma trận cánh tay "
              "được cập nhật.")
    bullet(d, "Trang số liệu: không đăng nhập bị chặn 401, tài khoản thường bị chặn, quản trị "
              "viên xem đủ 6 khối số liệu.")
    p(d, "Công tắc an toàn: đặt BANDIT_ENABLED=false hoặc AB_BANDIT_PCT=0 trong file cấu hình của "
         "dịch vụ Python là toàn bộ cơ chế khám phá tắt ngay, hệ thống quay về đúng hành vi "
         "Phase 2 — không cần sửa code hay triển khai lại.")
    p(d, "Lưu ý về dữ liệu: Phase 3 đang chạy trên dữ liệu mẫu được bơm sẵn để đạt ngưỡng "
         "(8 tuần × ~2.600 lượt hiển thị nguồn AI mỗi tuần, 230 lượt chuyển đổi neo vào đơn hàng "
         "thật) nhằm kiểm tra khả năng dự đoán trước. Khi chuyển sang dữ liệu thật, xoá dữ liệu "
         "mẫu bằng script dọn dẹp kèm manifest TRƯỚC khi bắt đầu tin các con số của bandit, vì "
         "trong thời gian thử nghiệm bandit học trên phản hồi của dữ liệu mẫu.")

    d.save(str(DOCX))
    print("Đã bổ sung PHẦN IV vào", DOCX)


if __name__ == "__main__":
    main()
