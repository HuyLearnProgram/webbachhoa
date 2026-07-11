# -*- coding: utf-8 -*-
r"""Sinh tài liệu thiết kế "Tìm kiếm thông minh" (Smart_Search_Roadmap.docx).

Tài liệu MỚI, tách riêng khỏi AI_Recommendation_Roadmap.docx (roadmap gợi ý đã đóng
sau Phase 3). Cùng bộ style: heading built-in, code-box Consolas nền xanh nhạt,
bảng "Light Grid Accent 1", sơ đồ matplotlib vẽ hộp bo góc.

Nội dung: thiết kế hệ tìm kiếm hybrid semantic + lexical + cá nhân hoá, phân kỳ
3 giai đoạn (A/B/C) — chốt cùng chủ dự án ngày 2026-07-11, CHƯA implement.

Chạy: cd recommendation-service && venv\Scripts\python scripts/docx_assets/build_smart_search_roadmap.py
(Đóng file Word trước khi chạy — script ghi đè trực tiếp Smart_Search_Roadmap.docx.)
"""
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor

ROOT = Path(__file__).resolve().parents[2]
DOCX = ROOT / "Smart_Search_Roadmap.docx"
ASSETS = Path(__file__).resolve().parent

BLUE, GREEN, ORANGE, RED, GRAY = "#2a78d6", "#1baf7a", "#eda100", "#e34948", "#5b5b5b"


# ---------------------------------------------------------------- docx helpers

def title(d, t):
    d.add_paragraph(t, style="Title")


def h1(d, t):
    d.add_paragraph(t, style="Heading 1")


def h2(d, t):
    d.add_paragraph(t, style="Heading 2")


def h3(d, t):
    d.add_paragraph(t, style="Heading 3")


def p(d, t):
    d.add_paragraph(t, style="Normal")


def bullet(d, t):
    d.add_paragraph(t, style="List Bullet")


def num(d, t):
    d.add_paragraph(t, style="List Number")


def picture(d, img: Path, caption: str):
    d.add_picture(str(img), width=Inches(6.3))
    d.add_paragraph(caption, style="Caption")


def code_box(d, text: str):
    """1x1 table nền xanh nhạt viền mảnh, font Consolas — đúng style code-box của
    AI_Recommendation_Roadmap.docx (nền F3F6F4, viền C7D3CB, chữ 2A3831 9pt)."""
    table = d.add_table(rows=1, cols=1)
    table.style = "Normal Table"
    cell = table.rows[0].cells[0]

    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F3F6F4")
    tcPr.append(shd)
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "C7D3CB")
        borders.append(el)
    tcPr.append(borders)

    para = cell.paragraphs[0]
    lines = text.strip("\n").split("\n")
    for i, line in enumerate(lines):
        run = para.add_run(line if i == 0 else "")
        if i > 0:
            run.add_break()
            run.text = line
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x2A, 0x38, 0x31)
    d.add_paragraph()


def data_table(d, headers, rows):
    table = d.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, htext in enumerate(headers):
        table.rows[0].cells[i].text = htext
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    d.add_paragraph()


# ---------------------------------------------------------------- sơ đồ

def _box(ax, xy, w, h, text, color, fs=10):
    x, y = xy
    ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.06",
                                fc=color, ec="none", alpha=0.14))
    ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.06",
                                fc="none", ec=color, lw=1.6))
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=fs,
            color="#222", wrap=True)


def _arrow(ax, a, b, color=GRAY, style="-|>", lw=1.6, ls="-"):
    ax.add_patch(FancyArrowPatch(a, b, arrowstyle=style, mutation_scale=16,
                                 color=color, lw=lw, linestyle=ls))


def draw_search_pipeline(path: Path) -> None:
    """Sơ đồ 1: luồng hybrid search FE → Java → Python với 3 tầng fallback."""
    fig, ax = plt.subplots(figsize=(10.5, 5.8))
    ax.set_xlim(0, 10.5)
    ax.set_ylim(0, 5.8)
    ax.axis("off")

    # hàng trên: luồng chính
    _box(ax, (0.2, 4.4), 2.0, 1.0, "Người dùng gõ\n\"đồ ăn ngon\"\n(FE gửi q qua JSON body)", BLUE, fs=9)
    _box(ax, (2.7, 4.4), 2.6, 1.0, "Java lọc điều kiện cứng ở DB\n→ allowed_ids (đúng filter)\n→ lexical_ids (khớp tên LIKE)", BLUE, fs=9)
    _box(ax, (5.8, 4.4), 2.4, 1.0, "Python /search/rank\nsemantic ∪ lexical\n+ cá nhân hoá + phổ biến", GREEN, fs=9)
    _box(ax, (8.6, 4.4), 1.7, 1.0, "Java phân trang\n+ hydrate DTO\n(giá luôn tươi)", BLUE, fs=9)
    _arrow(ax, (2.2, 4.9), (2.7, 4.9))
    _arrow(ax, (5.3, 4.9), (5.8, 4.9))
    _arrow(ax, (8.2, 4.9), (8.6, 4.9))

    # hàng giữa: 3 tầng fallback — mỗi tầng đặt thẳng dưới đúng tầng của nó,
    # mũi tên dọc "khi có sự cố" không cắt nhau
    _box(ax, (0.2, 2.3), 2.0, 1.0, "Tầng 3 — trong FE:\nendpoint mới lỗi/tắt cờ\n→ gọi lại GET\nproducts/search cũ", RED, fs=8.5)
    _box(ax, (2.7, 2.3), 2.6, 1.0, "Tầng 2 — trong Java:\nPython chết/chậm/rỗng\n→ LIKE như hiện tại (2.5s cắt)", RED, fs=9)
    _box(ax, (5.8, 2.3), 2.4, 1.0, "Tầng 1 — trong Python:\nencoder hỏng → tự hạ bậc\nONNX → TF-IDF → chỉ lexical", RED, fs=9)
    _arrow(ax, (1.2, 4.4), (1.2, 3.3), color=RED, ls=(0, (4, 3)))
    _arrow(ax, (4.0, 4.4), (4.0, 3.3), color=RED, ls=(0, (4, 3)))
    _arrow(ax, (7.0, 4.4), (7.0, 3.3), color=RED, ls=(0, (4, 3)))
    ax.text(1.35, 3.85, "sự cố", fontsize=8, color=RED, ha="left", style="italic")
    ax.text(4.15, 3.85, "sự cố", fontsize=8, color=RED, ha="left", style="italic")
    ax.text(7.15, 3.85, "sự cố", fontsize=8, color=RED, ha="left", style="italic")

    # nguyên tắc cấu trúc — đặt bên phải, cột không có tầng fallback
    ax.text(9.45, 2.8, "Python CHỈ được xếp hạng\ntrong allowed_ids —\nfilter được tôn trọng bằng\ncấu trúc, không bằng\nkỷ luật code",
            ha="center", fontsize=8.5, color=GREEN, style="italic")

    # hàng dưới
    _box(ax, (2.4, 0.5), 5.6, 1.0,
         "Không tầng nào được phép làm vỡ trang tìm kiếm —\ntệ nhất người dùng vẫn nhận kết quả LIKE y như hôm nay", ORANGE, fs=9.5)
    _arrow(ax, (1.2, 2.3), (3.6, 1.5), color=ORANGE, lw=1.1)
    _arrow(ax, (4.0, 2.3), (5.2, 1.5), color=ORANGE, lw=1.1)
    _arrow(ax, (7.0, 2.3), (6.8, 1.5), color=ORANGE, lw=1.1)

    fig.tight_layout()
    fig.savefig(path, dpi=160)
    plt.close(fig)


def draw_search_blend(path: Path) -> None:
    """Sơ đồ 2: cách trộn 4 nguồn điểm + gate ngưỡng, minh hoạ 2 kiểu query."""
    fig, ax = plt.subplots(figsize=(10.5, 5.6))
    ax.set_xlim(0, 10.5)
    ax.set_ylim(0, 5.6)
    ax.axis("off")

    # 4 nguồn điểm
    srcs = [
        ("Khớp tên (lexical)\nw = 1.0", "gõ đúng tên phải\nđứng #1", BLUE),
        ("Ngữ nghĩa (semantic)\nw = 0.8", "hiểu \"đồ ăn ngon\"\nqua embedding E5", GREEN),
        ("Cá nhân hoá\nw = 0.5", "tái dùng hồ sơ\n_profile_scores sẵn có", ORANGE),
        ("Phổ biến\nw = 0.15", "chỉ để phá hoà\n(tie-breaker)", GRAY),
    ]
    for i, (name, note, color) in enumerate(srcs):
        x = 0.3 + i * 2.6
        _box(ax, (x, 4.1), 2.3, 1.1, name, color, fs=9.5)
        ax.text(x + 1.15, 3.75, note, ha="center", fontsize=8, color=color)
        _arrow(ax, (x + 1.15, 3.5), (4.6 + (i - 1.5) * 0.5, 2.9), color=color, lw=1.2)

    _box(ax, (2.9, 1.9), 4.6, 1.0,
         "score = Σ wᵢ · điểmᵢ (chuẩn hoá min-max trong tập ứng viên)\nứng viên = lexical_ids ∪ semantic vượt ngưỡng", BLUE, fs=8.8)

    # gate ngưỡng
    _box(ax, (7.9, 1.9), 2.3, 1.0, "Gate chống query rác:\ntop-1 cosine < 0.82\n→ bỏ nhánh semantic", RED, fs=9)
    _arrow(ax, (7.9, 2.4), (7.55, 2.4), color=RED, ls=(0, (4, 3)))

    # 3 ví dụ minh hoạ
    ax.text(1.8, 1.15, "\"coca cola\" → lexical thắng:\nsản phẩm đúng tên đứng #1,\nsemantic chỉ bổ sung phía sau",
            ha="center", fontsize=8.5, color=BLUE, style="italic")
    ax.text(5.2, 1.15, "\"đồ ăn ngon\" → lexical rỗng:\nsemantic + cá nhân hoá quyết định,\nmón hợp gu khách lên đầu",
            ha="center", fontsize=8.5, color=GREEN, style="italic")
    ax.text(8.7, 1.15, "\"xzqw vbnk\" → cả 2 nhánh rỗng:\ntrả \"không tìm thấy\",\nkhông gợi bừa",
            ha="center", fontsize=8.5, color=RED, style="italic")
    _arrow(ax, (3.3, 1.9), (2.4, 1.55), color=BLUE, lw=1.0)
    _arrow(ax, (5.2, 1.9), (5.2, 1.55), color=GREEN, lw=1.0)
    _arrow(ax, (9.0, 1.9), (8.8, 1.55), color=RED, lw=1.0)

    ax.text(5.25, 0.3,
            "Khớp tên chính xác luôn thắng semantic; cá nhân hoá chỉ xếp lại thứ tự TRONG nhóm liên quan — không lật item khớp tên, không vượt filter.",
            ha="center", fontsize=9, color="#333", style="italic")

    fig.tight_layout()
    fig.savefig(path, dpi=160)
    plt.close(fig)


# ---------------------------------------------------------------- nội dung

def build(d):
    IMG_PIPE = ASSETS / "img_search_pipeline.png"
    IMG_BLEND = ASSETS / "img_search_blend.png"

    title(d, "Tìm kiếm thông minh (Smart Search)")
    p(d, "Thiết kế hệ tìm kiếm hybrid: hiểu nhu cầu tự do bằng embedding ngữ nghĩa, "
         "kết hợp khớp tên truyền thống và xếp hạng cá nhân hoá từ hệ gợi ý AI sẵn có. "
         "Tài liệu chốt ngày 11/07/2026. Trạng thái triển khai: Phase A ĐÃ XONG + verify "
         "(cùng ngày — chi tiết ở đầu mục Phase A), Phase B/C chưa làm. "
         "Tách riêng khỏi AI_Recommendation_Roadmap.docx (roadmap gợi ý đã hoàn tất Phase 0-3).")

    # ============================================================ PHẦN I
    h1(d, "PHẦN I — HIỆN TRẠNG & VẤN ĐỀ")

    h2(d, "1. Tìm kiếm hiện tại hoạt động thế nào")
    bullet(d, "Toàn bộ \"trí tuệ\" nằm ở frontend: buildProductNameFilter() (client/src/utils/helper.jsx) "
              "tách từ, bỏ dấu (stripDiacritics), dựng chuỗi RSQL productName~'từ1' and productName~'từ2'.")
    bullet(d, "Server (GET products/search, spring-filter) chỉ dịch RSQL → SQL LIKE %...% — không index "
              "(wildcard đầu chuỗi nên B-Tree cũng vô dụng), không full-text, không hiểu ngữ nghĩa.")
    bullet(d, "Sắp xếp: map thẳng vào Spring Pageable (tên/giá/rating/bán chạy/mới nhất). Không có khái niệm "
              "\"độ liên quan\" — mặc định trả theo thứ tự id.")
    bullet(d, "Product.description (MEDIUMTEXT) tồn tại nhưng chưa từng được search. Không có bảng tag/keyword nhu cầu.")
    bullet(d, "SearchLog (Phase 0 hệ gợi ý) đã ghi keyword + keyword_normalized + resultCount + clickedProductId "
              "kèm index idx_sl_keyword_norm, nhưng CHƯA có bất kỳ consumer nào đọc lại.")

    h2(d, "2. Vấn đề cần giải")
    num(d, "Người dùng tìm theo NHU CẦU chứ không chỉ tên: \"đồ ăn ngon\", \"thực phẩm sạch\", \"đồ sấy\"… "
           "— LIKE theo tên trả 0 kết quả với các query này.")
    num(d, "Kết quả cần xếp mặc định theo \"liên quan\" kiểu Shopee: trong phạm vi vẫn khớp từ khoá + bộ lọc, "
           "sản phẩm hợp sở thích cá nhân của khách (đã có sẵn trong hệ gợi ý AI) phải đứng đầu.")
    p(d, "Hướng đã chốt cùng chủ dự án: semantic/embedding thật (không dùng bảng từ khoá rule-based thủ công).")

    h2(d, "3. Ràng buộc kế thừa từ hệ thống")
    bullet(d, "Bug dấu tiếng Việt trong query param (root cause chưa rõ): mọi giá trị có dấu gửi qua query string "
              "đều có thể hỏng. Kinh nghiệm đã có: JSON body KHÔNG dính bug (các API tracking gửi tiếng Việt "
              "có dấu qua body thành công) → q bắt buộc đi qua POST body.")
    bullet(d, "Bean recommendationRestTemplate timeout connect 2000ms / read 2500ms — TUYỆT ĐỐI không nới "
              "(phá fail-fast fallback của Home/PDP/Cart). Search mới dùng chung bean, chung kỷ luật.")
    bullet(d, "Bài học thư viện Python trên Windows: implicit từng không build được (py3.12/numpy2) — mọi thư viện "
              "mới phải có spike script cài thử TRƯỚC khi thiết kế phụ thuộc vào nó.")
    bullet(d, "Rule bảo mật admin/public khai theo TỪNG HttpMethod — endpoint POST mới phải có rule riêng "
              "(bài học suýt lộ POST admin/recommendation-metrics/ab-pct).")
    bullet(d, "Filter cứng (active, quantity>0, category, giá, rating, khuyến mãi) phải được tôn trọng tuyệt đối — "
              "semantic chỉ quyết định THỨ TỰ và mở rộng recall bên trong phạm vi filter.")

    # ============================================================ PHẦN II
    h1(d, "PHẦN II — QUYẾT ĐỊNH KIẾN TRÚC")

    h2(d, "1. Model embedding: multilingual-e5-small qua ONNX Runtime")
    p(d, "Chọn intfloat/multilingual-e5-small (384 chiều, ~120MB): hỗ trợ tiếng Việt tốt (train trên kho đa ngữ "
         "có tiếng Việt), điểm retrieval đa ngữ cao hơn hẳn paraphrase-MiniLM cùng cỡ; all-MiniLM bị loại vì "
         "không hỗ trợ tiếng Việt. Repo HuggingFace của model có sẵn thư mục onnx/ — không cần PyTorch để export.")
    data_table(d,
        ["Bậc", "Phương án", "Điều kiện dùng"],
        [
            ["Plan A", "onnxruntime + tokenizers + model ONNX tải sẵn (~120MB)", "Mặc định — wheel Windows prebuilt, né PyTorch ~2.5GB"],
            ["Plan B", "sentence-transformers + torch CPU (cùng model e5-small)", "Nếu ONNX trục trặc khi spike"],
            ["Plan C", "TF-IDF query-vector (persist vectorizer của content_based.py vào ModelArtifacts — hiện đang bị vứt sau build)", "Backend degrade LUÔN tồn tại trong code, không chỉ là dự phòng cài đặt"],
        ])
    p(d, "Config search_embedding_backend: auto | onnx | torch | tfidf | off — auto thử lần lượt từng bậc, "
         "theo đúng pattern cf_backend đã dùng cho ALS.")
    p(d, "Chi tiết bắt buộc của họ E5 (quên là chất lượng rớt thảm mà KHÔNG báo lỗi): prefix \"query: \" cho câu "
         "truy vấn, \"passage: \" cho document sản phẩm; pooling = mean-pooling trên token embeddings rồi "
         "l2-normalize. Ghi thẳng vào docstring encoder + có test sanity cosine cố định trong pytest.")

    h2(d, "2. Hybrid recall — Java làm chủ filter và lexical, Python chỉ xếp hạng")
    picture(d, IMG_PIPE, "Sơ đồ 1 — Luồng tìm kiếm hybrid với 3 tầng fallback.")
    code_box(d, """FE ── POST /products/smart-search ──▶ Java SmartSearchService
     body JSON {q có dấu, sessionId}       1. allowed_ids = SELECT id theo filter cứng (spec)
     query param: filter RSQL (ASCII)      2. lexical_ids = SELECT id theo filter + LIKE từng từ của q
                                           3. POST Python /search/rank {q, allowed_ids, lexical_ids,
                                              user_id (từ JWT), session_id, limit}
                                           4. ranked list → phân trang → hydrate DTO (giữ thứ tự)
                                           5. Mọi lỗi/timeout/rỗng bất thường → fallback LIKE như cũ""")
    bullet(d, "Python CHỈ được chọn và xếp hạng trong allowed_ids — không bao giờ thêm id ngoài danh sách. "
              "Filter được tôn trọng bằng CẤU TRÚC, không phải bằng kỷ luật code.")
    bullet(d, "candidates = lexical_ids ∪ {p ∈ allowed_ids : semantic ≥ ngưỡng, top-K} — \"đồ ăn ngon\" "
              "(lexical rỗng) vẫn ra kết quả từ nhánh semantic; \"coca cola\" ra đúng sản phẩm tên khớp đứng #1.")
    bullet(d, "Sản phẩm tạo MỚI sau lần train gần nhất (chưa có trong ma trận embedding của Python) vẫn xuất hiện "
              "qua đường lexical — loại bỏ vấn đề catalog Python cũ tối đa 24h.")
    bullet(d, "q có dấu chỉ đi qua JSON body; chuỗi filter RSQL chỉ chứa id/số/enum ASCII nên vẫn đi query param, "
              "tái dùng binding @Filter sẵn có.")
    bullet(d, "Guard phía Java: recommendation.search.max-candidates=2000 — catalog vượt ngưỡng thì bỏ semantic, "
              "đi thẳng LIKE (bảo hiểm cho tương lai catalog phình; hiện tại vài trăm sản phẩm, list id vài KB).")

    h2(d, "3. Công thức trộn điểm")
    picture(d, IMG_BLEND, "Sơ đồ 2 — Trộn 4 nguồn điểm + gate chống query vô nghĩa.")
    code_box(d, """score(p) = w_lex·lex(p) + w_sem·sem_norm(p) + w_personal·personal_norm(p) + w_pop·pop_norm(p)

lex(p)      : 0 nếu p ∉ lexical_ids; ngược lại lexical_score() ∈ [1.0, 1.2]
              (đủ mọi token = 1.0, + 0.2 nếu khớp nguyên cụm)
sem(p)      : cosine(encode_query(q), emb[p]) — chỉ nhận nếu ≥ search_sem_min_sim
              VÀ top-1 toàn cục ≥ search_sem_top1_gate; chuẩn hoá min-max trong tập ứng viên
personal(p) : _profile_scores(...)[p] — tái dùng NGUYÊN TRẠNG app/main.py (~dòng 151-176),
              chạy cho cả user đăng nhập lẫn guest session; ngoài dict → 0
pop(p)      : tra art.popularity; chuẩn hoá trong tập ứng viên

Trọng số khởi điểm (override qua .env):
w_search_lexical=1.0 > w_search_semantic=0.8 > w_search_personal=0.5 > w_search_popularity=0.15""")
    bullet(d, "Triết lý trọng số: người gõ đúng tên sản phẩm phải thấy đúng nó đứng #1 (w_lex > w_sem); "
              "cá nhân hoá đủ mạnh để kéo sản phẩm hợp gu lên đầu TRONG nhóm liên quan nhưng không lật item "
              "khớp tên; phổ biến chỉ là tie-breaker.")
    bullet(d, "Search KHÔNG exclude seen — khác recommendation: người tìm kiếm muốn tìm lại cả thứ đã xem/mua. "
              "Khi tái dùng _profile_scores, bỏ qua giá trị seen trả về.")
    bullet(d, "Chuẩn hoá min-max trong phạm vi tập ứng viên, tái dùng _normalize của app/blend.py.")

    h2(d, "4. Ngưỡng chống query vô nghĩa")
    p(d, "Cosine của E5 dồn cụm cao (0.7–0.95 kể cả cặp không liên quan) nên không dùng 1 ngưỡng tuyệt đối đơn độc. "
         "Cơ chế 2 tầng, đều config qua .env:")
    data_table(d,
        ["Config", "Khởi điểm", "Vai trò"],
        [
            ["search_sem_top1_gate", "0.82", "Top-1 semantic trong allowed_ids < gate → semantic coi như KHÔNG có tín hiệu cho query này (query rác); nếu lexical cũng rỗng → matched=false, FE hiện empty-state"],
            ["search_sem_min_sim", "0.80", "Sàn tuyệt đối cho từng ứng viên lọt qua nhánh semantic (nâng từ 0.78 sau calibrate: vùng 0.78-0.81 là dải điểm của query rác)"],
            ["search_sem_top_k", "100", "Chặn đuôi số ứng viên semantic"],
            ["search_tfidf_min_sim", "0.05", "Bộ ngưỡng riêng khi backend=tfidf (thang điểm khác hẳn)"],
        ])
    p(d, "ĐÃ calibrate trên catalog thật 236 sản phẩm (scripts/calibrate_search.py, 11/07/2026): "
         "top-1 của 15 query nhu cầu thật nằm trong 0.836–0.894, của 5 query rác 0.785–0.809 — "
         "khoảng tách rõ, gate 0.82 nằm chính giữa (chặn nhầm 0/15 query thật, lọt 0/5 query rác). "
         "Sàn min_sim nâng 0.78 → 0.80 vì vùng 0.78–0.81 chính là dải điểm của query rác. "
         "Chất lượng định tính: \"đồ sấy\" ra đúng 10/10 Trái cây sấy, \"gia vị nấu ăn\"/\"đồ uống "
         "giải khát\" đúng nhóm. Hạn chế đã biết: query về ngành hàng catalog KHÔNG bán "
         "(VD \"đồ tẩy rửa\" trên catalog thuần thực phẩm) vẫn vượt gate với kết quả gần nghĩa "
         "nhất — gate không sửa được việc \"cửa hàng không bán thứ đó\". Khi catalog đổi lớn, "
         "chạy lại script calibrate rồi chỉnh .env.")

    h2(d, "5. Chuỗi fallback 3 tầng")
    data_table(d,
        ["Sự cố", "Hành vi"],
        [
            ["Encoder hỏng / model chưa tải", "Python tự hạ bậc backend (torch → tfidf → off), service vẫn lên; /search/rank vẫn phục vụ lexical + personal + popularity, algorithm_source=LEXICAL_ONLY"],
            ["Encode catalog lỗi giữa retrain", "Giữ emb_matrix cũ (carry-forward như CF), các artifact khác không ảnh hưởng"],
            ["Python chết / timeout / 500 / rỗng bất thường", "Java searchMode=LEXICAL_FALLBACK — kết quả y hệt hôm nay, cắt ở 2.5s"],
            ["matched=false (query rác, lexical cũng rỗng)", "Java trả trang rỗng NO_MATCH — KHÔNG gọi lại LIKE (lexical_ids đã rỗng từ DB, gọi lại chỉ phí)"],
            ["Endpoint smart-search 5xx / kill-switch off / lỗi mạng", "FE catch → gọi lại GET products/search cũ với buildProductNameFilter"],
            ["User chọn sort tường minh (giá/rating…)", "Vẫn recall thông minh, sort DB-side trong tập matched: WHERE id IN (rankedIds) ORDER BY <field>"],
        ])
    p(d, "Ngân sách độ trễ thực tế: encode 1 query ONNX CPU ~10–50ms (có LRU cache 512 theo query chuẩn hoá) + "
         "cosine brute-force vài trăm sản phẩm <1ms + 2–3 query hồ sơ ~10–30ms — dư dả trong budget 2.5s. "
         "Catalog nhỏ nên brute-force numpy là đủ, CHƯA cần FAISS/vector index (tránh over-engineer sớm).")

    # ============================================================ PHẦN III
    h1(d, "PHẦN III — PHÂN KỲ TRIỂN KHAI")

    # ---------------- Phase A
    h2(d, "Phase A — Nền tảng semantic phía Python (chưa đụng Java/FE) — ĐÃ XONG 11/07/2026")
    p(d, "TRẠNG THÁI: đã code xong + verify đầy đủ ngày 11/07/2026. Spike PASS ngay Plan A "
         "(onnxruntime 1.27 + tokenizers + huggingface-hub đều wheel prebuilt trên Windows/py3.12/"
         "numpy2 — không lặp lại số phận implicit; model ONNX 470MB, lưu ý con số ~120MB ở Phần II "
         "là ước lượng ban đầu chưa chính xác). Pytest 15/15; calibrate trên catalog thật 236 sản "
         "phẩm (kết quả ở Phần II mục 4); E2E qua uvicorn thật đủ 5 kịch bản: semantic rescue "
         "\"đồ sấy\" trả đúng 10/10 Trái cây sấy (đối chiếu DB), query rác matched=false, "
         "allowed_ids không rò rỉ, lexical thắng semantic, session lạ không crash. "
         "Service KHÔNG tự tải model lúc start — máy mới chạy scripts/check_embedding_env.py 1 lần.")
    p(d, "Mục tiêu: Python service encode được catalog + query tiếng Việt, expose POST /search/rank hoàn chỉnh "
         "(semantic ∪ lexical + cá nhân hoá + phổ biến), verify bằng curl/pytest. Chưa user nào nhìn thấy — "
         "zero rủi ro UX. Không có ngưỡng dữ liệu (semantic không cần interaction).")
    h3(d, "A.1. Tiên quyết: spike cài đặt PASS trước khi viết code chính")
    bullet(d, "scripts/check_embedding_env.py: cài thử onnxruntime + tokenizers + huggingface-hub vào venv, "
              "tải model về data/models/multilingual-e5-small/, encode 3 câu tiếng Việt, in cosine "
              "\"đồ sấy khô\" vs \"mít sấy giòn\" vs \"nước rửa chén\". Pass mới commit requirements.")
    bullet(d, "Fail → thử Plan B (torch CPU); fail nữa → chốt Plan C (TF-IDF) và ghi rõ vào README.")
    h3(d, "A.2. File mới — package app/search/")
    data_table(d,
        ["File", "Nội dung chính"],
        [
            ["app/search/encoder.py", "Singleton load 1 lần lúc process start (KHÔNG load lại mỗi retrain). encode_passages(texts) batch prefix \"passage: \"; encode_query(q) prefix \"query: \" + LRU cache 512; backend() → onnx|torch|tfidf|off, auto rơi bậc khi import/model lỗi, log WARNING rõ"],
            ["app/search/lexical.py", "Chuẩn hoá không dấu + lowercase (port stripDiacritics/normalizeKeyword); lexical_score(query_tokens, name_norm) ∈ [1.0, 1.2] — chỉ TINH CHỈNH thứ tự trong lexical_ids Java gửi sang, membership do Java quyết"],
            ["app/search/embeddings.py", "build(df): document = \"passage: {tên}. {category}. {description strip HTML, cắt 500 ký tự}\"; cache encode theo (product_id, md5(document)) vào data/emb_cache.npz — retrain nightly chỉ encode sản phẩm mới/sửa"],
            ["app/search/rank.py", "rank_search(q, allowed_ids, lexical_ids, user_id, session_id, art, limit) — hợp nhất ứng viên, gate ngưỡng, blend 4 nguồn, trả ranked list"],
        ])
    h3(d, "A.3. File sửa")
    bullet(d, "app/store.py — ModelArtifacts thêm: emb_matrix (np [N,384] l2-normalized, hàng khớp product_ids, "
              "dùng chung pid_to_idx), name_norm (dict id → tên chuẩn hoá), tfidf_vectorizer (Plan C).")
    bullet(d, "app/candidates/content_based.py — trả thêm vectorizer + name_norm (df đã có sẵn trong tay).")
    bullet(d, "app/trainer.py — build embedding trong train_all(); lỗi encode → emb_matrix=None + log, "
              "KHÔNG giết các artifact khác (giống pattern _train_cf).")
    bullet(d, "app/main.py — endpoint POST /search/rank; /health thêm search_backend, emb_products_indexed.")
    bullet(d, "app/config.py — nhóm search_* và w_search_* (bảng ở Phần II); requirements.txt theo kết quả spike.")
    h3(d, "A.4. API contract POST /search/rank")
    code_box(d, """// Request
{ "query": "đồ ăn ngon", "allowed_ids": [1,5,9,...], "lexical_ids": [5],
  "user_id": 12, "session_id": "abc", "limit": 300 }

// Response 200
{ "request_id": "…", "model_version": "…",
  "algorithm_source": "SEMANTIC_HYBRID",   // TFIDF_HYBRID khi degrade, LEXICAL_ONLY khi backend off
  "matched": true,
  "items": [ { "product_id": 5, "score": 1.83, "source": "LEXICAL" } ] }
  // source ∈ LEXICAL | SEMANTIC | PERSONALIZED — nguồn đóng góp lớn nhất, cùng convention blend.py
  // matched=false, items=[] khi cả 2 nhánh rỗng; 503 khi model chưa train (Java tự fallback)""")
    h3(d, "A.5. Verify & rủi ro")
    bullet(d, "pytest tests/test_search_rank.py: ứng viên ngoài allowed_ids không bao giờ lọt; gate chặn query rác; "
              "sanity cosine cố định (cặp gần nghĩa > cặp xa nghĩa — bắt lỗi quên prefix/pooling sai); "
              ".env override trọng số có tác dụng; backend=off vẫn trả lexical.")
    bullet(d, "scripts/calibrate_search.py: bảng top-10 + cosine cho ~20 query thật → chỉnh 2 ngưỡng.")
    bullet(d, "Rủi ro: wheel lỗi (đã có bậc thang A/B/C); quên prefix E5 (test sanity); ngưỡng sai trên seed ít ỏi "
              "(calibrate + .env); description seed kém kéo lệch embedding → nếu calibrate thấy nhiễu, "
              "hạ description khỏi document (chỉ tên + category) — quyết định tại bước calibrate.")

    # ---------------- Phase B
    h2(d, "Phase B — Tích hợp Java + FE: relevance mặc định, cá nhân hoá thật")
    p(d, "Mục tiêu: user thật gõ \"đồ ăn ngon\" ra kết quả; mặc định xếp \"Liên quan\" có cá nhân hoá; filter cứng "
         "giữ nguyên hiệu lực; sort tường minh hoạt động trong phạm vi recall mới; Python chết → search y như cũ. "
         "Tiên quyết: Phase A pass calibrate.")
    h3(d, "B.1. Java")
    bullet(d, "MỚI service/SmartSearchService.java — inject bean @Qualifier(\"recommendationRestTemplate\") "
              "(2000/2500ms, không nới), luồng 5 bước như Sơ đồ 1; hydrate slice trang (10 id/lần) bằng "
              "findActiveDtoByIdIn giữ nguyên thứ tự (pattern hydrate của RecommendationService); "
              "userId CHỈ derive từ JWT (getCurrentUserOrNull), không nhận từ client.")
    bullet(d, "MỚI domain/request/SmartSearchRequestDTO.java {q, sessionId}; "
              "domain/response/recommendation/PythonSearchRankResponse.java.")
    bullet(d, "SỬA ProductController.java — POST products/smart-search?filter=&page=&size=&sort= (body {q, sessionId}); "
              "bind @Filter Specification<Product> từ query param như cũ (filter chỉ ASCII).")
    bullet(d, "SỬA ProductRepository — projection id-only theo Specification (cạnh searchProduct sẵn có).")
    bullet(d, "SỬA SecurityConfiguration — đúng 1 rule permitAll cho ĐÚNG HttpMethod.POST + ĐÚNG path "
              "products/smart-search (bài học per-method; không wildcard mở nhầm các POST products/* của admin).")
    bullet(d, "SỬA application.properties/yml — recommendation.search.enabled (kill-switch: off → controller đi thẳng "
              "fallback LIKE), max-candidates=2000, rank-limit=300.")
    bullet(d, "Trả SmartSearchResultDTO = PaginationDTO (meta/result y hệt cũ — FE tái dùng toàn bộ render) "
              "+ searchMode (SEMANTIC_HYBRID | TFIDF_HYBRID | LEXICAL_FALLBACK | NO_MATCH) + requestId.")
    h3(d, "B.2. Frontend")
    bullet(d, "apis/product.js — apiSmartSearch({q, filter, page, size, sort}); sessionId lấy từ "
              "getOrCreateSessionId() của apis/recommendation.js.")
    bullet(d, "utils/constants.js — sortProductOption thêm option đầu {value:'', text:'Liên quan'} — mặc định khi "
              "có search term.")
    bullet(d, "pages/guest/Product.jsx — effect fetch rẽ nhánh: có searchTerm → apiSmartSearch với filter = các mảnh "
              "hiện tại TRỪ buildProductNameFilter (active/category/promotion/rating/price giữ nguyên); "
              "catch/lỗi → chạy lại đúng nhánh cũ (fallback tầng FE). Không có searchTerm → đường cũ 100%.")
    bullet(d, "SearchBar.jsx autocomplete GIỮ NGUYÊN LIKE — autocomplete cần <100ms và match prefix tên là đúng UX; "
              "semantic để dành cho trang kết quả.")
    bullet(d, "Known-limitation chấp nhận ở B: apiGetMaxPrice (slider giá) vẫn tính theo LIKE → khoảng giá có thể "
              "hẹp hơn tập semantic; vá ở C nếu thấy cấn.")
    h3(d, "B.3. Verify E2E (Selenium, theo skill verify của dự án)")
    num(d, "Search \"đồ sấy\" → ≥1 kết quả (semantic rescue; hôm nay = 0).")
    num(d, "Search đúng tên 1 sản phẩm seed → sản phẩm đó đứng #1 (lexical thắng semantic).")
    num(d, "Cá nhân hoá: session A xem 3 sản phẩm category X → search từ khoá chung 2 category → item X đứng trên; "
           "so sánh profile Chrome sạch thứ tự khác.")
    num(d, "Filter cứng: search + tick category + kéo giá → 100% kết quả đúng category/giá.")
    num(d, "Kill-Python drill: tắt uvicorn → search vẫn trả kết quả LIKE, không lỗi UI.")
    num(d, "Query rác \"xzqw vbnk\" → empty-state \"Không tìm thấy\", không crash.")
    num(d, "Phân trang relevance: page 1 ∪ page 2 không trùng id.")
    num(d, "Sort \"Giá tăng dần\" khi đang search semantic → kết quả trong tập liên quan, đúng thứ tự giá.")
    num(d, "Regression: \"sữa\" và \"sua\" ra cùng kết quả.")
    h3(d, "B.4. Rủi ro")
    bullet(d, "Mở nhầm quyền POST (đã có tiền sự) → rule per-method + test 401 các POST products/* khác.")
    bullet(d, "@Filter binding trên POST khác biệt tinh vi → dự phòng: nhận filter string trong body, convert bằng "
              "FilterSpecificationConverter của spring-filter.")
    bullet(d, "Test cũ assume thứ tự id → rà Selenium/TestCase liên quan search trước khi merge.")
    bullet(d, "Personal score decay theo giờ → thứ tự \"nhảy nhẹ\" giữa 2 lần load: chấp nhận (drift chậm).")

    # ---------------- Phase C
    h2(d, "Phase C — Đo lường, gợi ý từ khoá, tinh chỉnh")
    p(d, "Tiên quyết: Phase B chạy ~1–2 tuần lấy log (trên seed thì chủ động tạo traffic qua Selenium script).")
    h3(d, "C.1. Migration + đo lường (thuần Java — số liệu phải sống kể cả khi Python chết)")
    code_box(d, """ALTER TABLE search_logs ADD COLUMN search_mode VARCHAR(30) NULL;
ALTER TABLE search_logs ADD COLUMN clicked_position INT NULL;
-- nullable: client cũ chưa gửi field mới cũng không vỡ""")
    bullet(d, "FE: apiTrackSearch gửi thêm searchMode từ response smart-search; handleSearchResultClick gửi vị trí "
              "item = index + (page-1)*size.")
    bullet(d, "Endpoint admin MỚI GET admin/search-metrics?days= (rule ADMIN per-method): "
              "Search CTR theo mode (so LEXICAL_FALLBACK vs SEMANTIC_HYBRID trực tiếp); Zero-result rate + top "
              "keyword 0 kết quả (mỏ vàng nhập hàng/đặt tên lại); Semantic-rescue rate (mode semantic có kết quả "
              "trong khi lexical_ids rỗng — Java truyền cờ lexicalEmpty vào log); Avg click position theo mode.")
    bullet(d, "FE: section \"Chất lượng tìm kiếm\" trong dashboard Gợi ý AI (RecommendationMetrics.jsx) — 4 stat + "
              "bảng top zero-result keywords, kèm giải thích trong recommendationMetricsExplain.jsx, hiện n (cỡ mẫu) "
              "+ cảnh báo \"chưa đủ dữ liệu\" như pattern AB-pct.")
    h3(d, "C.2. Autocomplete \"Mọi người cũng tìm\" từ SearchLog")
    bullet(d, "Java GET products/search-suggestions?prefix=&limit=8 (permitAll GET; prefix đã strip dấu client-side "
              "→ match keyword_normalized ăn index idx_sl_keyword_norm — né bug dấu bằng thiết kế).")
    bullet(d, "Query: GROUP BY keyword_normalized, HAVING COUNT(*) >= 2 AND MAX(result_count) > 0 — không gợi ý "
              "keyword từng 0 kết quả, không gợi ý keyword đơn lẻ (vừa chất lượng vừa tránh lộ query cá biệt của "
              "người khác); trả bản CÓ DẤU (keyword) xuất hiện nhiều nhất; cache in-memory 5 phút.")
    bullet(d, "FE SearchBar.jsx: khối \"Mọi người cũng tìm\" phía trên autocomplete sản phẩm, debounce dùng chung; "
              "click → navigate /products?search=. Lỗi/rỗng → SearchBar hoạt động như cũ (catch nuốt lỗi).")
    h3(d, "C.3. Tinh chỉnh")
    bullet(d, "Chạy lại calibrate_search.py với top-50 keyword thật từ search_logs → chỉnh search_sem_* và "
              "w_search_* qua .env.")
    bullet(d, "Tuỳ chọn nếu rảnh: cho chỉnh 2 trọng số từ dashboard admin qua pattern /internal/ab-config + "
              "_update_env_file sẵn có.")

    # ============================================================ PHẦN IV
    h1(d, "PHẦN IV — DANH MỤC FILE & CẤU HÌNH TỔNG HỢP")

    h2(d, "1. Danh mục file theo tầng")
    data_table(d,
        ["Tầng", "File mới", "File sửa"],
        [
            ["Python (recommendation-service/)",
             "app/search/{encoder, lexical, rank, embeddings}.py; scripts/check_embedding_env.py; scripts/calibrate_search.py; tests/test_search_rank.py",
             "app/config.py; app/store.py; app/trainer.py; app/main.py; app/candidates/content_based.py; requirements.txt"],
            ["Java (server/.../webnongsan/)",
             "service/SmartSearchService.java; domain/request/SmartSearchRequestDTO.java; domain/response/recommendation/PythonSearchRankResponse.java",
             "controller/ProductController.java; repository/ProductRepository.java; config/SecurityConfiguration.java; application.properties (Phase C: service+controller EventTracking)"],
            ["Frontend (client/src/)",
             "—",
             "apis/product.js; pages/guest/Product.jsx; utils/constants.js (Phase C: components/SearchBar.jsx; pages/admin/RecommendationMetrics.jsx; apis/recommendation.js)"],
            ["Khác",
             "SQL/ migration search_logs (Phase C); Selenium/TestCase bộ test smart-search",
             "—"],
        ])

    h2(d, "2. Config mới")
    code_box(d, """# ===== Python recommendation-service (.env / app/config.py) =====
search_embedding_backend=auto        # auto | onnx | torch | tfidf | off
search_model_dir=data/models/multilingual-e5-small
search_sem_min_sim=0.80
search_sem_top1_gate=0.82
search_tfidf_min_sim=0.05
search_sem_top_k=100
search_max_ranked=500                # trần số item trả cho Java phân trang
w_search_lexical=1.0
w_search_semantic=0.8
w_search_personal=0.5
w_search_popularity=0.15

# ===== Java (application.properties) =====
recommendation.search.enabled=true   # kill-switch không cần sửa code
recommendation.search.max-candidates=2000
recommendation.search.rank-limit=300""")

    h2(d, "3. Nguyên tắc bất di bất dịch (nhắc lại)")
    num(d, "Không tầng nào được làm vỡ trang tìm kiếm — tệ nhất là kết quả LIKE như hôm nay.")
    num(d, "Filter cứng tôn trọng tuyệt đối bằng cấu trúc (Python chỉ xếp hạng trong allowed_ids).")
    num(d, "Tiếng Việt có dấu chỉ đi qua JSON body; không nới timeout bean recommendationRestTemplate.")
    num(d, "Spike cài thư viện trên Windows PASS rồi mới code; rule bảo mật khai đúng từng HttpMethod.")
    num(d, "Mọi ngưỡng/trọng số qua .env, có bước calibrate bằng dữ liệu thật trước khi tin.")


def main() -> None:
    draw_search_pipeline(ASSETS / "img_search_pipeline.png")
    draw_search_blend(ASSETS / "img_search_blend.png")

    d = Document()
    build(d)
    d.save(str(DOCX))
    print(f"OK: {DOCX}")


if __name__ == "__main__":
    main()
